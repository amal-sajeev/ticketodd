# AI-Powered Grievance Redressal System
# FastAPI + MongoDB + Qdrant + OpenAI

import os
import re
import uuid
import asyncio
import hashlib
import json
import logging
import math
import statistics
import difflib
import ipaddress
from pathlib import Path
from datetime import datetime, timedelta, timezone
from typing import List, Optional, Dict, Any
from enum import Enum
from contextlib import asynccontextmanager
from concurrent.futures import ThreadPoolExecutor
import numpy as np

FACE_MATCH_THRESHOLD = 0.5  # Strict threshold for high security
import base64
import io

import httpx
import openai as openai_mod
import uvicorn
import gridfs
from bson import ObjectId
from fastapi import FastAPI, HTTPException, BackgroundTasks, Depends, status, Query, File, UploadFile, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import OAuth2PasswordBearer
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from fastapi.responses import HTMLResponse, RedirectResponse, StreamingResponse
from starlette.requests import Request
from pydantic import BaseModel, Field, field_validator, model_validator
from pymongo import MongoClient, ReturnDocument, GEOSPHERE
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, VectorParams, PointStruct, Filter, FieldCondition, MatchValue, PointIdsList
from openai import AsyncOpenAI
from jose import JWTError, jwt
from passlib.context import CryptContext
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------
from dotenv import load_dotenv
# Try multiple .env locations: next to this file, one level up, then cwd
_script_dir = Path(__file__).resolve().parent
_env_candidates = [
    _script_dir / ".env",            # demo/.env
    _script_dir.parent / ".env",     # ticketer-main/.env
    Path.cwd() / ".env",             # current working directory
]
_env_loaded = False
for _env_path in _env_candidates:
    if _env_path.is_file():
        load_dotenv(_env_path, override=True)
        _env_loaded = True
        break
if not _env_loaded:
    load_dotenv(override=True)  # fall back to python-dotenv's own search

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

BASE_DIR = Path(__file__).resolve().parent
MONGODB_URL = os.getenv("MONGODB_URL", "mongodb://localhost:27017")
QDRANT_URL = os.getenv("QDRANT_URL", "http://localhost:6333")
QDRANT_API_KEY = os.getenv("QDRANT_API_KEY")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-5-mini")
EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "text-embedding-3-large")
JWT_SECRET = os.getenv("JWT_SECRET", "")
if not JWT_SECRET or len(JWT_SECRET) < 32:
    raise RuntimeError(
        "FATAL: JWT_SECRET must be set in the environment and be at least 32 characters. "
        "Generate one with: python -c \"import secrets; print(secrets.token_urlsafe(48))\""
    )
JWT_ALGORITHM = "HS256"
JWT_EXPIRE_HOURS = 24
EMBEDDING_DIM = 3072  # text-embedding-3-large

OPENWEATHERMAP_API_KEY = os.getenv("OPENWEATHERMAP_API_KEY", "")

openai_client = AsyncOpenAI(api_key=OPENAI_API_KEY)
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="/auth/login", auto_error=False)

# GridFS instance (initialized in startup_db)
gfs = None

# Census data for impact scoring (loaded at startup)
CENSUS_DATA: Dict[str, Any] = {}

# Load census data
_census_path = BASE_DIR / "static" / "odisha_census_data.json"
if _census_path.is_file():
    with open(_census_path, "r", encoding="utf-8") as _cf:
        CENSUS_DATA = json.load(_cf)

# Allowed upload MIME types and limits
ALLOWED_MIME_TYPES = {
    "image/jpeg", "image/png", "image/webp", "image/gif",
    "audio/mpeg", "audio/wav", "audio/ogg", "audio/mp4",
    "video/mp4", "video/webm", "video/quicktime",
}
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
MAX_FILES_PER_GRIEVANCE = 5

# Scheme document upload limits
SCHEME_ALLOWED_MIMES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "image/jpeg", "image/png",
}
SCHEME_MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
SCHEME_MAX_FILES = 3
OPENAI_VISION_MODEL = os.getenv("OPENAI_VISION_MODEL", "gpt-4o")

# Department survival necessity scores for impact scoring
DEPT_SURVIVAL_SCORES = {
    "rural_water_supply": 25, "sanitation": 20, "rural_housing": 18,
    "mgnregs": 15, "infrastructure": 12, "rural_livelihoods": 10,
    "panchayati_raj": 8, "general": 5,
}

# ---------------------------------------------------------------------------
# Enums
# ---------------------------------------------------------------------------
class Department(str, Enum):
    PANCHAYATI_RAJ = "panchayati_raj"
    RURAL_WATER_SUPPLY = "rural_water_supply"
    MGNREGS = "mgnregs"
    RURAL_HOUSING = "rural_housing"
    RURAL_LIVELIHOODS = "rural_livelihoods"
    SANITATION = "sanitation"
    INFRASTRUCTURE = "infrastructure"
    GENERAL = "general"

class OfficerCategory(str, Enum):
    BLOCK_DEV_OFFICER = "block_dev_officer"
    DISTRICT_PANCHAYAT_OFFICER = "district_panchayat_officer"
    EXECUTIVE_ENGINEER_RWSS = "executive_engineer_rwss"
    DRDA_PROJECT_DIRECTOR = "drda_project_director"
    GP_SECRETARY = "gp_secretary"
    MGNREGS_PROGRAMME_OFFICER = "mgnregs_programme_officer"
    GENERAL_OFFICER = "general_officer"
    SENIOR_OFFICER = "senior_officer"

class Priority(str, Enum):
    LOW = "low"
    MEDIUM = "medium"
    HIGH = "high"
    URGENT = "urgent"

class GrievanceStatus(str, Enum):
    PENDING = "pending"
    IN_PROGRESS = "in_progress"
    RESOLVED = "resolved"
    ESCALATED = "escalated"

class Sentiment(str, Enum):
    POSITIVE = "positive"
    NEUTRAL = "neutral"
    NEGATIVE = "negative"
    FRUSTRATED = "frustrated"

class ResolutionType(str, Enum):
    AI = "ai"
    MANUAL = "manual"
    HYBRID = "hybrid"

class ResolutionTier(str, Enum):
    SELF_RESOLVABLE = "self_resolvable"
    OFFICER_ACTION = "officer_action"
    ESCALATION = "escalation"

class NoteType(str, Enum):
    INTERNAL = "internal"
    CITIZEN_FACING = "citizen_facing"

class Language(str, Enum):
    ODIA = "odia"
    HINDI = "hindi"
    ENGLISH = "english"

class UserRole(str, Enum):
    CITIZEN = "citizen"
    OFFICER = "officer"
    ADMIN = "admin"

# ---------------------------------------------------------------------------
# Pydantic Models
# ---------------------------------------------------------------------------
class Geolocation(BaseModel):
    latitude: Optional[float] = Field(None, ge=-90, le=90)
    longitude: Optional[float] = Field(None, ge=-180, le=180)
    type: Optional[str] = None
    coordinates: Optional[List[float]] = None

    @field_validator('coordinates')
    @classmethod
    def validate_coordinates(cls, v):
        if v is not None and len(v) != 2:
            raise ValueError("Coordinates must be [longitude, latitude]")
        return v

    @model_validator(mode='before')
    @classmethod
    def check_location_format(cls, values):
        if isinstance(values, dict):
            if values.get('coordinates') is not None:
                if len(values['coordinates']) != 2:
                    raise ValueError("Coordinates must be [longitude, latitude]")
                values['longitude'] = values['coordinates'][0]
                values['latitude'] = values['coordinates'][1]
            elif values.get('latitude') is not None and values.get('longitude') is not None:
                values['coordinates'] = [values['longitude'], values['latitude']]
                values['type'] = 'Point'
        return values

class UserCreate(BaseModel):
    username: str = Field(..., min_length=3, max_length=50)
    password: str = Field(..., min_length=6, max_length=72)
    full_name: str = Field(..., max_length=200)
    email: str = Field(..., max_length=320)
    phone: Optional[str] = Field(None, max_length=20)
    role: UserRole = UserRole.CITIZEN
    department: Optional[Department] = None
    face_descriptor: Optional[List[float]] = None

class UserUpdate(BaseModel):
    full_name: Optional[str] = Field(None, max_length=200)
    email: Optional[str] = Field(None, max_length=320)
    phone: Optional[str] = Field(None, max_length=20)
    role: Optional[UserRole] = None
    department: Optional[Department] = None
    password: Optional[str] = Field(None, min_length=6, max_length=72)
    face_descriptor: Optional[List[float]] = None

class UserLogin(BaseModel):
    username: Optional[str] = None
    password: Optional[str] = None
    face_descriptor: Optional[List[float]] = None

class UserResponse(BaseModel):
    id: str
    username: str
    full_name: str
    email: str
    phone: Optional[str] = None
    role: UserRole
    department: Optional[str] = None
    created_at: datetime
    has_face_id: bool = False

class TokenResponse(BaseModel):
    access_token: str
    token_type: str = "bearer"
    user: UserResponse

class GrievanceCreate(BaseModel):
    title: str = Field(..., max_length=200)
    description: str = Field(..., max_length=5000)
    citizen_name: Optional[str] = None
    citizen_email: Optional[str] = None
    citizen_phone: Optional[str] = None
    is_anonymous: bool = False
    is_public: bool = False
    language: Language = Language.ENGLISH
    district: Optional[str] = None
    location: Optional[Geolocation] = None

class GrievanceResponse(BaseModel):
    id: str
    tracking_number: str
    title: str
    description: str
    citizen_name: Optional[str] = None
    citizen_email: Optional[str] = None
    citizen_phone: Optional[str] = None
    is_anonymous: bool = False
    is_public: bool = False
    language: Language
    district: Optional[str] = None
    department: Department
    priority: Priority
    officer_category: OfficerCategory
    status: GrievanceStatus
    sentiment: Sentiment
    created_at: datetime
    updated_at: datetime
    sla_deadline: Optional[datetime] = None
    estimated_resolution_deadline: Optional[datetime] = None
    ai_resolution: Optional[str] = None
    manual_resolution: Optional[str] = None
    resolution_type: Optional[ResolutionType] = None
    confidence_score: Optional[float] = None
    resolution_tier: Optional[ResolutionTier] = None
    assigned_officer: Optional[str] = None
    resolution_feedback: Optional[int] = Field(None, ge=1, le=5)
    notes: List[Dict] = Field(default_factory=list)
    location: Optional[Geolocation] = None
    citizen_user_id: Optional[str] = None
    attachments: List[Dict] = Field(default_factory=list)
    impact_score: Optional[int] = None
    scheme_match: Optional[Dict] = None
    sub_tasks: List[Dict] = Field(default_factory=list)
    systemic_issue_id: Optional[str] = None

class GrievanceTrackResponse(BaseModel):
    id: str
    tracking_number: str
    title: str
    status: GrievanceStatus
    department: Department
    priority: Priority
    created_at: datetime
    updated_at: datetime
    sla_deadline: Optional[datetime] = None
    ai_resolution: Optional[str] = None
    manual_resolution: Optional[str] = None
    resolution_type: Optional[ResolutionType] = None
    resolution_tier: Optional[ResolutionTier] = None

class ChatMessage(BaseModel):
    message: str = Field(..., max_length=2000)
    language: Language = Language.ENGLISH
    conversation_history: List[Dict[str, str]] = Field(default_factory=list)

class ChatResponse(BaseModel):
    reply: str
    sources: List[Dict[str, Any]] = Field(default_factory=list)
    filed_grievance: Optional[Dict[str, Any]] = None

class KnowledgeEntry(BaseModel):
    title: str = Field(..., max_length=500)
    content: str = Field(..., max_length=20000)
    category: Department

class ServiceMemoryEntry(BaseModel):
    query: str = Field(..., max_length=5000)
    resolution: str = Field(..., max_length=10000)
    category: Department
    agent_name: str = Field(..., max_length=200)

class SchemeEntry(BaseModel):
    name: str = Field(..., max_length=500)
    description: str = Field(..., max_length=10000)
    eligibility: str = Field(..., max_length=5000)
    department: Department
    how_to_apply: str = Field(..., max_length=5000)
    eligibility_questions: List[Dict[str, str]] = Field(default_factory=list)

class SchemeResponse(BaseModel):
    id: str; name: str; description: str; eligibility: str
    department: str; how_to_apply: str; created_at: str
    eligibility_questions: List[Dict[str, str]] = Field(default_factory=list)
    documents: List[Dict[str, Any]] = Field(default_factory=list)

class GenerateQuestionsRequest(BaseModel):
    eligibility: str
    scheme_name: str = ""

class UpdateQuestionsRequest(BaseModel):
    questions: List[Dict[str, str]]

class DocumentationResponse(BaseModel):
    id: str; title: str; content: str; category: str; created_at: str

class ServiceMemoryResponse(BaseModel):
    id: str; query: str; resolution: str; category: str; agent_name: str; created_at: str
    tracking_number: Optional[str] = None
    grievance_id: Optional[str] = None

class ManualResolution(BaseModel):
    resolution: str = Field(..., max_length=10000)
    officer: str = Field(..., max_length=200)
    add_to_service_memory: bool = True

class ResolutionFeedback(BaseModel):
    rating: int = Field(..., ge=1, le=5)
    comments: Optional[str] = Field(None, max_length=2000)

class GrievanceNote(BaseModel):
    content: str = Field(..., max_length=5000)
    officer: str = Field(..., max_length=200)
    note_type: NoteType = NoteType.INTERNAL

class GrievanceAssignment(BaseModel):
    officer_name: str = Field(..., max_length=200)

class PublicGrievanceResponse(BaseModel):
    id: str
    title: str
    description_snippet: str
    department: str
    priority: str
    district: Optional[str] = None
    location: Optional[Dict] = None
    created_at: datetime
    vouch_count: int = 0
    evidence_count: int = 0
    distance_km: Optional[float] = None

class VouchRequest(BaseModel):
    comment: Optional[str] = Field(None, max_length=500)

class SystemicIssueResponse(BaseModel):
    id: str
    title: str
    root_cause_analysis: str
    department: str
    district: Optional[str] = None
    affected_area_center: Optional[Dict] = None
    affected_radius_km: Optional[float] = None
    grievance_ids: List[str] = Field(default_factory=list)
    estimated_population_affected: int = 0
    priority: str
    status: str
    created_at: datetime
    updated_at: datetime
    assigned_officer: Optional[str] = None

class ForecastResponse(BaseModel):
    id: str
    forecast_date: datetime
    forecast_period_start: datetime
    forecast_period_end: datetime
    predictions: List[Dict] = Field(default_factory=list)
    generated_by: str = "system"

class AnalyticsResponse(BaseModel):
    total_grievances: int
    self_resolved: int = 0
    ai_drafted: int = 0
    escalated_to_human: int
    avg_resolution_time: float
    pending_count: int = 0
    sla_breached_count: int = 0
    deadline_alerts: int = 0
    status_distribution: Dict[str, int] = Field(default_factory=dict)
    sentiment_distribution: Dict[str, int]
    department_distribution: Dict[str, int]
    top_districts: List[Dict[str, Any]]
    top_complaints: List[Dict[str, Any]]

class GeoAnalyticsResponse(BaseModel):
    hotspot_coordinates: List[Dict[str, float]]
    region_distribution: Dict[str, int]
    district_distribution: Dict[str, int]

# ---------------------------------------------------------------------------
# App & Globals
# ---------------------------------------------------------------------------
app = FastAPI(title="Panchayati Raj & Drinking Water Department — Grievance Portal")
limiter = Limiter(key_func=get_remote_address)
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

# ---------------------------------------------------------------------------
# Security Headers Middleware
# ---------------------------------------------------------------------------
from starlette.middleware.base import BaseHTTPMiddleware

class SecurityHeadersMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request, call_next):
        response = await call_next(request)
        response.headers["X-Content-Type-Options"] = "nosniff"
        response.headers["X-Frame-Options"] = "SAMEORIGIN"
        response.headers["X-XSS-Protection"] = "1; mode=block"
        response.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"
        response.headers["Permissions-Policy"] = "geolocation=(self), camera=(), microphone=()"
        response.headers["Content-Security-Policy"] = (
            "default-src 'self'; "
            "script-src 'self' 'unsafe-inline' https://cdn.jsdelivr.net; "
            "style-src 'self' 'unsafe-inline' https://fonts.googleapis.com; "
            "img-src 'self' data: blob:; "
            "media-src 'self' blob:; "
            "font-src 'self' https://fonts.gstatic.com; "
            "connect-src 'self' https://cdn.jsdelivr.net; "
            "frame-src 'self'; "
            "frame-ancestors 'self'"
        )
        return response

app.add_middleware(SecurityHeadersMiddleware)
db_client = None
db = None
qdrant = None
gfs = None
executor = ThreadPoolExecutor(max_workers=10)

from jinja2 import Environment, FileSystemLoader, select_autoescape
_jinja_env = Environment(
    loader=FileSystemLoader(str(BASE_DIR / "templates")),
    autoescape=select_autoescape(["html", "htm", "xml"]),
)
templates = Jinja2Templates(directory=str(BASE_DIR / "templates"))
templates.env = _jinja_env
app.mount("/static", StaticFiles(directory=str(BASE_DIR / "static")), name="static")

# ---------------------------------------------------------------------------
# Startup / Shutdown
# ---------------------------------------------------------------------------
@asynccontextmanager
async def lifespan(app: FastAPI):
    await startup_db()
    await startup_qdrant()
    logger.info("OpenAI model: %s | Embedding: %s", OPENAI_MODEL, EMBEDDING_MODEL)
    yield
    if db_client:
        db_client.close()

app.router.lifespan_context = lifespan

async def startup_db():
    global db_client, db, gfs
    db_client = MongoClient(MONGODB_URL)
    db = db_client.grievance_system
    gfs = gridfs.GridFS(db)
    loop = asyncio.get_event_loop()
    await loop.run_in_executor(executor, db.grievances.create_index, "created_at")
    await loop.run_in_executor(executor, db.grievances.create_index, "status")
    await loop.run_in_executor(executor, db.grievances.create_index, "department")
    await loop.run_in_executor(executor, db.grievances.create_index, "priority")
    await loop.run_in_executor(executor, db.grievances.create_index, "tracking_number")
    await loop.run_in_executor(executor, lambda: db.users.create_index([("username", 1)], unique=True))
    # 2dsphere index for geospatial queries on public feed
    await loop.run_in_executor(executor, lambda: db.grievances.create_index([("location", GEOSPHERE)]))
    await loop.run_in_executor(executor, lambda: db.grievances.create_index("citizen_user_id"))
    await loop.run_in_executor(executor, lambda: db.grievances.create_index("is_public"))
    # Spam tracking indexes
    await loop.run_in_executor(executor, lambda: db.spam_tracking.create_index("_id"))
    # Vouches indexes
    await loop.run_in_executor(executor, lambda: db.vouches.create_index("grievance_id"))
    await loop.run_in_executor(executor, lambda: db.vouches.create_index([("grievance_id", 1), ("user_id", 1)], unique=True))
    # Systemic issues
    await loop.run_in_executor(executor, lambda: db.systemic_issues.create_index("status"))
    await loop.run_in_executor(executor, lambda: db.systemic_issues.create_index("department"))
    # Officer analytics
    await loop.run_in_executor(executor, lambda: db.officer_analytics.create_index("_id"))
    # Forecasts
    await loop.run_in_executor(executor, lambda: db.forecasts.create_index([("forecast_date", -1)]))
    logger.info("Database initialized with all indexes")

async def startup_qdrant():
    global qdrant
    try:
        if QDRANT_API_KEY:
            qdrant = QdrantClient(url=QDRANT_URL, api_key=QDRANT_API_KEY, timeout=30)
        else:
            qdrant = QdrantClient(url=QDRANT_URL)
        qdrant.get_collections()
    except Exception as e:
        logger.warning("Qdrant unavailable (%s) — knowledge base features disabled.", e)
        qdrant = None
        return
    for collection in ["documentation", "service_memory", "schemes"]:
        try:
            existing = qdrant.get_collections()
            names = [c.name for c in existing.collections]
            if collection not in names:
                qdrant.create_collection(
                    collection_name=collection,
                    vectors_config=VectorParams(size=EMBEDDING_DIM, distance=Distance.COSINE))
                logger.info("Created collection: %s", collection)
        except Exception as e:
            logger.error("Qdrant collection %s error: %s", collection, e)

# ---------------------------------------------------------------------------
# Dependencies
# ---------------------------------------------------------------------------
async def get_db():
    return db

async def get_qdrant():
    return qdrant

# ---------------------------------------------------------------------------
# Auth Helpers
# ---------------------------------------------------------------------------
import bcrypt

# ... (rest of imports)

# Remove pwd_context = CryptContext(...) line or just ignore it
# We will replace the functions that use it

def hash_password(password: str) -> str:
    if len(password.encode("utf-8")) > 72:
        raise ValueError("Password cannot exceed 72 bytes")
    # return pwd_context.hash(password)
    return bcrypt.hashpw(password.encode("utf-8"), bcrypt.gensalt()).decode("utf-8")

def verify_password(plain: str, hashed: str) -> bool:
    # return pwd_context.verify(plain, hashed)
    return bcrypt.checkpw(plain.encode("utf-8"), hashed.encode("utf-8"))

def create_access_token(data: dict) -> str:
    to_encode = data.copy()
    to_encode["exp"] = datetime.now(timezone.utc) + timedelta(hours=JWT_EXPIRE_HOURS)
    return jwt.encode(to_encode, JWT_SECRET, algorithm=JWT_ALGORITHM)

_token_blacklist: set = set()

async def get_current_user(token: Optional[str] = Depends(oauth2_scheme), db=Depends(get_db)):
    if token is None:
        raise HTTPException(status_code=401, detail="Not authenticated")
    if token in _token_blacklist:
        raise HTTPException(status_code=401, detail="Token has been revoked")
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        username: str = payload.get("sub")
        if username is None:
            raise HTTPException(status_code=401, detail="Invalid token")
    except JWTError:
        raise HTTPException(status_code=401, detail="Invalid token")
    loop = asyncio.get_event_loop()
    user = await loop.run_in_executor(executor, db.users.find_one, {"username": username})
    if user is None:
        raise HTTPException(status_code=401, detail="User not found")
    return user

async def get_optional_user(token: Optional[str] = Depends(oauth2_scheme), db=Depends(get_db)):
    if token is None:
        return None
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        username = payload.get("sub")
        if username is None:
            return None
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(executor, db.users.find_one, {"username": username})
    except JWTError:
        return None

def require_role(*roles):
    async def role_checker(user=Depends(get_current_user)):
        if user["role"] not in roles:
            raise HTTPException(status_code=403, detail="Insufficient permissions")
        return user
    return role_checker

def user_to_response(user: dict) -> UserResponse:
    return UserResponse(
        id=str(user["_id"]), username=user["username"], full_name=user["full_name"],
        email=user["email"], phone=user.get("phone"), role=user["role"],
        department=user.get("department"), created_at=user["created_at"],
        has_face_id=bool(user.get("face_descriptor")))

# ---------------------------------------------------------------------------
# AI Helpers: Cache, Retry, Truncation
# ---------------------------------------------------------------------------
_embedding_cache: Dict[str, List[float]] = {}

def truncate_text(text: str, max_chars: int = 3000) -> str:
    if len(text) <= max_chars:
        return text
    return text[:max_chars] + "..."

async def get_openai_embedding(text: str) -> List[float]:
    text = truncate_text(text, 2000)
    key = hashlib.sha256(text.encode()).hexdigest()
    if key in _embedding_cache:
        return _embedding_cache[key]
    response = await openai_client.embeddings.create(model=EMBEDDING_MODEL, input=text)
    vec = response.data[0].embedding
    _embedding_cache[key] = vec
    if len(_embedding_cache) > 500:
        oldest = next(iter(_embedding_cache))
        del _embedding_cache[oldest]
    return vec

async def openai_chat(messages: list, json_mode: bool = False, max_retries: int = 3,
                      **extra_kwargs) -> Optional[str]:
    kwargs = {}
    if json_mode:
        kwargs["response_format"] = {"type": "json_object"}
    for attempt in range(max_retries):
        try:
            resp = await openai_client.chat.completions.create(
                model=OPENAI_MODEL, messages=messages, **kwargs)
            return resp.choices[0].message.content.strip()
        except (openai_mod.RateLimitError, openai_mod.APIConnectionError) as e:
            logger.warning("OpenAI retry %d: %s", attempt + 1, e)
            if attempt == max_retries - 1:
                raise
            await asyncio.sleep(2 ** attempt)
        except Exception as e:
            logger.error("OpenAI error: %s", e)
            return None
    return None

# ---------------------------------------------------------------------------
# Utility Helpers
# ---------------------------------------------------------------------------
def generate_tracking_number(db) -> str:
    import secrets
    import string
    counter = db.counters.find_one_and_update(
        {"_id": "grievance"}, {"$inc": {"seq": 1}},
        upsert=True, return_document=ReturnDocument.AFTER)
    # Append a random alphanumeric suffix to prevent enumeration
    random_suffix = ''.join(secrets.choice(string.ascii_uppercase + string.digits) for _ in range(6))
    return f"GRV-{datetime.now(timezone.utc).year}-{counter['seq']:06d}{random_suffix}"

def calculate_sla_deadline(priority: Priority) -> datetime:
    hours_map = {Priority.URGENT: 24, Priority.HIGH: 72, Priority.MEDIUM: 168, Priority.LOW: 360}
    return datetime.now(timezone.utc) + timedelta(hours=hours_map.get(priority, 168))

def convert_db_grievance(g: dict) -> dict:
    if g.get("location"):
        g["location"] = {
            "type": g["location"].get("type"),
            "coordinates": g["location"].get("coordinates"),
            "longitude": g["location"]["coordinates"][0] if g["location"].get("coordinates") else None,
            "latitude": g["location"]["coordinates"][1] if g["location"].get("coordinates") else None,
        }
    return g

LANGUAGE_NAMES = {"odia": "Odia (ଓଡ଼ିଆ)", "hindi": "Hindi (हिन्दी)", "english": "English"}

# ---------------------------------------------------------------------------
# Input Sanitization Helpers
# ---------------------------------------------------------------------------
ODISHA_DISTRICTS = [
    "Angul","Balangir","Balasore","Bargarh","Bhadrak","Boudh","Cuttack","Deogarh",
    "Dhenkanal","Gajapati","Ganjam","Jagatsinghpur","Jajpur","Jharsuguda","Kalahandi",
    "Kandhamal","Kendrapara","Kendujhar","Khordha","Koraput","Malkangiri","Mayurbhanj",
    "Nabarangpur","Nayagarh","Nuapada","Puri","Rayagada","Sambalpur","Subarnapur","Sundargarh"
]

def sanitize_str(value: str) -> str:
    """Ensure a value is a plain string, not a dict/list that could be a NoSQL operator."""
    if not isinstance(value, str):
        raise HTTPException(status_code=400, detail="Invalid parameter type")
    return str(value)

def validate_uuid(value: str, param_name: str = "id") -> str:
    """Validate that a string is a valid UUID format."""
    value = sanitize_str(value)
    try:
        uuid.UUID(value)
    except (ValueError, AttributeError):
        raise HTTPException(status_code=400, detail=f"Invalid {param_name} format")
    return value


# ---------------------------------------------------------------------------
# Face Recognition Helpers
# ---------------------------------------------------------------------------
def face_distance(face1: List[float], face2: List[float]) -> float:
    """Calculate Euclidean distance between two face descriptors"""
    if not face1 or not face2:
        return float('inf')
    return float(np.linalg.norm(np.array(face1) - np.array(face2)))

async def find_user_by_face(descriptor: List[float], db) -> Optional[dict]:
    """Find a user with a matching face descriptor."""
    if not descriptor:
        return None
    loop = asyncio.get_event_loop()
    # Fetch all users with face descriptors (optimize in prod with vector DB if needed)
    users_with_faces = await loop.run_in_executor(
        executor, lambda: list(db.users.find({"face_descriptor": {"$exists": True}})))
    
    best_match = None
    best_distance = float('inf')
    
    for user in users_with_faces:
        stored_face = user.get("face_descriptor")
        if not stored_face:
            continue
        dist = face_distance(descriptor, stored_face)
        if dist < best_distance:
            best_distance = dist
            best_match = user
            
    if best_match and best_distance < FACE_MATCH_THRESHOLD:
        return best_match
    return None

# ---------------------------------------------------------------------------
# AI Service Classes
# ---------------------------------------------------------------------------
class SentimentAnalyzer:
    @staticmethod
    async def analyze_sentiment(text: str) -> Sentiment:
        text = truncate_text(text)
        prompt = (
            "Analyze the sentiment of this citizen grievance about government services. "
            "Classify as exactly one of: positive, neutral, negative, frustrated.\n\n"
            f'Grievance: "{text}"\n\nRespond with only one word.'
        )
        try:
            result = await openai_chat([{"role": "user", "content": prompt}])
            if result:
                r = result.lower()
                if "frustrated" in r: return Sentiment.FRUSTRATED
                if "negative" in r: return Sentiment.NEGATIVE
                if "positive" in r: return Sentiment.POSITIVE
            return Sentiment.NEUTRAL
        except Exception as e:
            logger.error("Sentiment analysis error: %s", e)
            return Sentiment.NEUTRAL

class GrievanceClassifier:
    @staticmethod
    async def classify(title: str, description: str) -> tuple:
        title = truncate_text(title, 200)
        description = truncate_text(description, 2800)
        prompt = (
            "Classify this citizen grievance for the Odisha Panchayati Raj & Drinking Water Department Grievance Portal.\n\n"
            f'Title: "{title}"\nDescription: "{description}"\n\n'
            "Return a JSON object with exactly these keys:\n"
            '- "department": one of [panchayati_raj, rural_water_supply, mgnregs, rural_housing, rural_livelihoods, sanitation, infrastructure, general]\n'
            '- "priority": one of [low, medium, high, urgent]\n'
            '- "officer": one of [block_dev_officer, district_panchayat_officer, executive_engineer_rwss, drda_project_director, gp_secretary, mgnregs_programme_officer, general_officer, senior_officer]\n'
            '- "tier": one of [self_resolvable, officer_action, escalation]\n\n'
            "Department guide: panchayati_raj=GP/Block/Zilla Parishad governance/elections/meetings/Sarpanch issues, "
            "rural_water_supply=Jal Jeevan Mission/Basudha/RWSS/bore wells/pipelines/water quality, "
            "mgnregs=MGNREGA job cards/wage payments/worksite complaints, "
            "rural_housing=PMAY-Gramin/Biju Pucca Ghar/Nirman Shramik housing, "
            "rural_livelihoods=NRLM/OLM/SHG loans/Mission Shakti/skill development, "
            "sanitation=Swachh Bharat Mission Gramin/ODF/toilet construction/SLWM, "
            "infrastructure=BGBO scheme/rural roads/bridges/community buildings/Finance Commission grants, "
            "general=other.\n"
            "Priority guide: urgent=life-threatening/water contamination, high=significant hardship, medium=standard, low=info request.\n"
            "Tier guide: self_resolvable=citizen can fix it themselves with correct info or guidance "
            "(application status queries, document requirements, portal/login issues, process questions, "
            "eligibility queries), officer_action=needs physical inspection, repair, dispatch, approval, "
            "or any government action, escalation=sensitive/corruption/fund misuse/repeated complaints/"
            "life-threatening situations."
        )
        try:
            result = await openai_chat([{"role": "user", "content": prompt}], json_mode=True)
            if result:
                data = json.loads(result)
                dept = Department(data["department"]) if data.get("department") in [e.value for e in Department] else Department.GENERAL
                pri = Priority(data["priority"]) if data.get("priority") in [e.value for e in Priority] else Priority.MEDIUM
                off = OfficerCategory(data["officer"]) if data.get("officer") in [e.value for e in OfficerCategory] else OfficerCategory.GENERAL_OFFICER
                tier = ResolutionTier(data["tier"]) if data.get("tier") in [e.value for e in ResolutionTier] else ResolutionTier.OFFICER_ACTION
                return dept, pri, off, tier
        except Exception as e:
            logger.error("Classification error: %s", e)
        return Department.GENERAL, Priority.MEDIUM, OfficerCategory.GENERAL_OFFICER, ResolutionTier.OFFICER_ACTION

class KnowledgeSearcher:
    @staticmethod
    async def search_documentation(query: str, limit: int = 5) -> List[Dict]:
        try:
            query_vector = await get_openai_embedding(query)
            if not qdrant: return []
            results = qdrant.query_points(collection_name="documentation", query=query_vector,
                                    limit=limit, score_threshold=0.4, with_payload=True, with_vectors=False)
            return [{"content": p.payload.get("content", ""), "title": p.payload.get("title", ""),
                      "category": p.payload.get("category", ""), "score": p.score} for p in results.points]
        except Exception as e:
            logger.error("Documentation search error: %s", e)
            return []

    @staticmethod
    async def search_service_memory(query: str, limit: int = 5) -> List[Dict]:
        try:
            query_vector = await get_openai_embedding(query)
            if not qdrant: return []
            results = qdrant.query_points(collection_name="service_memory", query=query_vector,
                                    limit=limit, score_threshold=0.45, with_payload=True, with_vectors=False)
            return [{"query": p.payload.get("query", ""), "resolution": p.payload.get("resolution", ""),
                      "category": p.payload.get("category", ""), "score": p.score} for p in results.points]
        except Exception as e:
            logger.error("Service memory search error: %s", e)
            return []

    @staticmethod
    async def search_schemes(query: str, limit: int = 5) -> List[Dict]:
        try:
            query_vector = await get_openai_embedding(query)
            if not qdrant: return []
            results = qdrant.query_points(collection_name="schemes", query=query_vector,
                                    limit=limit, score_threshold=0.35, with_payload=True, with_vectors=False)
            return [{"id": str(p.id), "name": p.payload.get("name", ""), "description": p.payload.get("description", ""),
                      "eligibility": p.payload.get("eligibility", ""),
                      "how_to_apply": p.payload.get("how_to_apply", ""),
                      "department": p.payload.get("department", ""),
                      "created_at": p.payload.get("created_at", ""),
                      "eligibility_questions": p.payload.get("eligibility_questions", []),
                      "score": p.score} for p in results.points]
        except Exception as e:
            logger.error("Schemes search error: %s", e)
            return []

class ResolutionGenerator:
    @staticmethod
    def _build_context(context_sources: List[Dict]) -> str:
        if not context_sources:
            return ""
        text = "\n\nRelevant knowledge:\n"
        for source in context_sources[:3]:
            if source.get("resolution"):
                text += f"- Previous resolution: {truncate_text(source['resolution'], 500)}\n"
            elif source.get("content"):
                text += f"- Documentation: {truncate_text(source['content'], 500)}\n"
            elif source.get("how_to_apply"):
                text += f"- Scheme: {source['name']} - {truncate_text(source['description'], 300)}\n"
        return text

    @staticmethod
    def _calc_confidence(context_sources: List[Dict], text: str) -> float:
        confidence = 0.3
        if context_sources:
            confidence += max(s.get("score", 0) for s in context_sources) * 0.4
        if text and len(text) > 50:
            confidence += 0.2
        return min(confidence, 0.95)

    @staticmethod
    async def generate_citizen_guidance(title: str, description: str, citizen_name: str,
                                         sentiment: str, language: str, context_sources: List[Dict]) -> tuple:
        """Generate self-help guidance addressed directly to the citizen."""
        context_text = ResolutionGenerator._build_context(context_sources)
        lang_name = LANGUAGE_NAMES.get(language, "English")
        prompt = (
            f"You are Seva, the AI assistant built into the Odisha Panchayati Raj & Drinking Water Department Grievance Portal.\n"
            f"Respond in {lang_name}. Use Markdown formatting.\n\n"
            "You are providing a ONE-TIME guidance message to a citizen who just submitted a grievance through this portal. "
            "This is NOT a conversation — the citizen cannot reply to you. Your response must be complete and self-contained.\n\n"
            "CRITICAL RULES:\n"
            "- Prioritize DIGITAL SELF-SERVICE: government portals, websites, online tracking, toll-free helplines, and mobile apps.\n"
            "- This portal IS the citizen's point of contact with the government. Do NOT tell them to 'visit the GP office' or "
            "'contact the BDO' as a first step — they filed here precisely to avoid that.\n"
            "- Only suggest visiting a physical office as a LAST RESORT if no digital option exists for the specific issue.\n"
            "- Mention specific portal URLs (e.g. awaassoft.nic.in, egramswaraj.gov.in) and toll-free helplines (e.g. 1916, 1800-11-6446) where applicable.\n"
            "- NEVER claim any government action has been taken.\n"
            "- NEVER invent reference numbers, case IDs, or tracking codes.\n"
            "- NEVER write as if you can respond to follow-up questions (no 'tell me more', 'let me know', 'I can help you with...').\n"
            "- Use direct, professional tone.\n\n"
            "FORMATTING RULES:\n"
            "- Keep the total response under 150 words.\n"
            "- Structure: one brief line acknowledging the issue, then a '### What You Can Do' section with 2-4 "
            "numbered steps (online actions first, physical visits only if unavoidable), then one closing line "
            "noting that their grievance has also been recorded in the system for follow-up.\n"
            "- Use **bold** for key terms, portal names, and helpline numbers.\n"
            "- Do not repeat the grievance description back.\n\n"
            f"**Issue:** {truncate_text(title, 200)} — {truncate_text(description, 1500)}\n"
            f"{context_text}\n"
            f"Citizen: {citizen_name or 'Citizen'} | Sentiment: {sentiment}"
        )
        try:
            text = await openai_chat([{"role": "user", "content": prompt}])
            if not text:
                return "Your grievance has been noted and will be reviewed by a government officer.", 0.0
            return text, ResolutionGenerator._calc_confidence(context_sources, text)
        except Exception as e:
            logger.error("Citizen guidance generation error: %s", e)
            return "Your grievance has been noted and will be reviewed by a government officer.", 0.0

    @staticmethod
    async def generate_officer_draft(title: str, description: str, citizen_name: str,
                                      sentiment: str, language: str, context_sources: List[Dict]) -> tuple:
        """Generate a draft response for an officer to review before sending."""
        context_text = ResolutionGenerator._build_context(context_sources)
        lang_name = LANGUAGE_NAMES.get(language, "English")
        prompt = (
            f"You are an AI assistant drafting a response for an officer in the Odisha "
            f"Panchayati Raj & Drinking Water Department. The officer will review and edit this before sending.\n"
            f"Draft in {lang_name}. Use Markdown formatting.\n\n"
            "CRITICAL RULES:\n"
            "- This is a DRAFT for officer review — not a final response.\n"
            "- Suggest what actions the department should take.\n"
            "- NEVER claim any action has already been taken.\n"
            "- NEVER invent reference numbers, case IDs, or tracking codes.\n"
            "- Use future tense: 'We will dispatch...', 'The department will investigate...'.\n"
            "- Only mention specific departments/agencies if clearly relevant.\n\n"
            "FORMATTING RULES:\n"
            "- Keep the total response under 150 words.\n"
            "- Structure: one line acknowledging receipt, then a '### Recommended Actions' section with "
            "2-4 bullet points for the officer to act on, then one closing line.\n"
            "- Use **bold** for key terms and deadlines.\n"
            "- Be direct and professional. Do not repeat the grievance description back.\n\n"
            f"**Grievance:** {truncate_text(title, 200)} — {truncate_text(description, 1500)}\n"
            f"{context_text}\n"
            f"Citizen: {citizen_name or 'Anonymous'} | Sentiment: {sentiment}"
        )
        try:
            text = await openai_chat([{"role": "user", "content": prompt}])
            if not text:
                return "Your grievance has been noted and will be reviewed by a government officer.", 0.0
            return text, ResolutionGenerator._calc_confidence(context_sources, text)
        except Exception as e:
            logger.error("Officer draft generation error: %s", e)
            return "Your grievance has been noted and will be reviewed by a government officer.", 0.0

# ---------------------------------------------------------------------------
# Eligibility Question Generator
# ---------------------------------------------------------------------------
async def generate_eligibility_questions(eligibility_text: str, scheme_name: str) -> list:
    prompt = (
        f'For the government scheme "{scheme_name}", the eligibility criteria are:\n'
        f'"{truncate_text(eligibility_text, 1500)}"\n\n'
        'Generate 3-5 simple yes/no questions that determine if a citizen is eligible. '
        'Each question should be easy for a common citizen to understand. Return JSON: '
        '{"questions": [{"question": "...", "eligible_answer": "yes"}, ...]}'
    )
    try:
        result = await openai_chat([{"role": "user", "content": prompt}], json_mode=True)
        if not result:
            return []
        data = json.loads(result)
        questions = data.get("questions", data) if isinstance(data, dict) else data
        if isinstance(questions, list):
            return [{"question": q.get("question", ""), "eligible_answer": q.get("eligible_answer", "yes")}
                    for q in questions if q.get("question")]
        return []
    except Exception as e:
        logger.error("Eligibility question generation error: %s", e)
        return []

# ---------------------------------------------------------------------------
# Spam Detection Engine
# ---------------------------------------------------------------------------
class SpamDetector:
    @staticmethod
    async def check_and_update(user_id: str, title: str, description: str, ip: str) -> tuple:
        """Returns (is_blocked: bool, reason: str). Updates spam_tracking."""
        loop = asyncio.get_event_loop()
        record = await loop.run_in_executor(executor, db.spam_tracking.find_one, {"_id": user_id})
        if not record:
            record = {
                "_id": user_id, "filing_timestamps": [], "duplicate_hashes": [],
                "spam_score": 0.0, "is_blocked": False, "blocked_at": None,
                "photo_id_file_id": None, "photo_id_status": "none",
                "admin_notified": False, "pattern_flags": [], "ip_addresses": [],
            }
        if record.get("is_blocked") and record.get("photo_id_status") != "approved":
            return True, "Account blocked due to spam activity. Please upload a photo ID to verify your identity."
        now = datetime.now(timezone.utc)
        content_hash = hashlib.sha256(f"{title}|{description}".lower().strip().encode()).hexdigest()
        timestamps = record.get("filing_timestamps", [])
        # Ensure timestamps are tz-aware (MongoDB may return naive datetimes)
        timestamps = [t.replace(tzinfo=timezone.utc) if t.tzinfo is None else t for t in timestamps]
        timestamps = [t for t in timestamps if (now - t).total_seconds() < 3600]
        timestamps.append(now)
        dup_hashes = record.get("duplicate_hashes", [])
        # Ensure dup_hash timestamps are tz-aware
        for h in dup_hashes:
            if h["ts"].tzinfo is None:
                h["ts"] = h["ts"].replace(tzinfo=timezone.utc)
        dup_hashes = [h for h in dup_hashes if (now - h["ts"]).total_seconds() < 86400]
        score = record.get("spam_score", 0.0)
        reason_parts = []
        # Frequency check: >5 in 1 hour
        if len(timestamps) > 5:
            score = min(score + 0.25, 1.0)
            reason_parts.append("high filing frequency")
        # Duplicate content check
        if any(h["hash"] == content_hash for h in dup_hashes):
            score = min(score + 0.3, 1.0)
            reason_parts.append("duplicate content detected")
        dup_hashes.append({"hash": content_hash, "ts": now})
        # Short description spam
        if len(description.strip()) < 20 and len(timestamps) > 2:
            score = min(score + 0.15, 1.0)
            reason_parts.append("very short descriptions")
        ip_list = record.get("ip_addresses", [])
        if ip and ip not in ip_list:
            ip_list.append(ip)
        is_blocked = score >= 0.7
        update = {
            "$set": {
                "filing_timestamps": timestamps, "duplicate_hashes": dup_hashes,
                "spam_score": round(score, 3), "is_blocked": is_blocked,
                "ip_addresses": ip_list[-20:],
            }
        }
        if is_blocked and not record.get("blocked_at"):
            update["$set"]["blocked_at"] = now
        await loop.run_in_executor(executor, lambda: db.spam_tracking.update_one(
            {"_id": user_id}, update, upsert=True))
        if is_blocked:
            return True, "Account blocked: " + ", ".join(reason_parts)
        return False, ""

# ---------------------------------------------------------------------------
# Passive Pattern Analyzer
# ---------------------------------------------------------------------------
class PatternAnalyzer:
    @staticmethod
    async def analyze(user_id: str, title: str, description: str, ip: str):
        """Background task: analyze filing patterns and update spam score."""
        try:
            loop = asyncio.get_event_loop()
            text = f"{title} {description}"
            flags = []
            # Content similarity with user's recent filings
            recent = await loop.run_in_executor(executor, lambda: list(
                db.grievances.find({"citizen_user_id": user_id}).sort("created_at", -1).limit(10)))
            if len(recent) >= 3 and qdrant:
                try:
                    new_emb = await get_openai_embedding(text)
                    high_sim_count = 0
                    for g in recent[1:]:
                        old_emb = await get_openai_embedding(f"{g['title']} {g['description']}")
                        sim = sum(a * b for a, b in zip(new_emb, old_emb)) / (
                            (sum(a**2 for a in new_emb)**0.5) * (sum(b**2 for b in old_emb)**0.5) + 1e-10)
                        if sim > 0.92:
                            high_sim_count += 1
                    if high_sim_count >= 2:
                        flags.append({"type": "content_similarity", "detail": f"{high_sim_count} highly similar filings",
                                      "timestamp": datetime.now(timezone.utc)})
                except Exception:
                    pass
            # Keyword quality check
            words = description.split()
            if words:
                caps_ratio = sum(1 for w in words if w.isupper()) / len(words)
                if caps_ratio > 0.5 and len(words) > 5:
                    flags.append({"type": "keyword_stuffing", "detail": f"{caps_ratio:.0%} uppercase words",
                                  "timestamp": datetime.now(timezone.utc)})
            # Cross-user duplicate detection
            content_hash = hashlib.sha256(f"{title}|{description}".lower().strip().encode()).hexdigest()
            dup_count = await loop.run_in_executor(executor, lambda: db.spam_tracking.count_documents({
                "_id": {"$ne": user_id},
                "duplicate_hashes.hash": content_hash}))
            if dup_count > 0:
                flags.append({"type": "cross_user_duplicate",
                              "detail": f"Same content from {dup_count} other account(s)",
                              "timestamp": datetime.now(timezone.utc)})
            # IP correlation
            if ip:
                ip_users = await loop.run_in_executor(executor, lambda: db.spam_tracking.count_documents({
                    "_id": {"$ne": user_id}, "ip_addresses": ip}))
                if ip_users >= 2:
                    flags.append({"type": "ip_correlation",
                                  "detail": f"IP shared with {ip_users} other accounts",
                                  "timestamp": datetime.now(timezone.utc)})
            if flags:
                # Compute weighted score adjustment
                score_add = len(flags) * 0.1
                await loop.run_in_executor(executor, lambda: db.spam_tracking.update_one(
                    {"_id": user_id},
                    {"$push": {"pattern_flags": {"$each": flags, "$slice": -50}},
                     "$inc": {"spam_score": min(score_add, 0.3)}},
                    upsert=True))
        except Exception as e:
            logger.error("Pattern analysis error: %s", e)

# ---------------------------------------------------------------------------
# Officer Anomaly Detector
# ---------------------------------------------------------------------------
class OfficerAnomalyDetector:
    @staticmethod
    async def check_after_resolution(officer_id: str, officer_name: str, resolution_text: str):
        """Background task: check for anomalous officer resolution patterns."""
        try:
            loop = asyncio.get_event_loop()
            now = datetime.now(timezone.utc)
            today_str = now.strftime("%Y-%m-%d")
            record = await loop.run_in_executor(executor, db.officer_analytics.find_one, {"_id": officer_id})
            if not record:
                record = {"_id": officer_id, "daily_resolution_counts": {},
                          "avg_resolution_text_length": 0.0, "avg_time_to_resolve_hours": 0.0,
                          "priority_distribution": {}, "anomaly_flags": [], "baseline_computed_at": now}
            daily = record.get("daily_resolution_counts", {})
            daily[today_str] = daily.get(today_str, 0) + 1
            # Keep only last 30 days
            cutoff = (now - timedelta(days=30)).strftime("%Y-%m-%d")
            daily = {k: v for k, v in daily.items() if k >= cutoff}
            new_flags = []
            # Bulk rubber-stamping check
            today_count = daily.get(today_str, 0)
            if today_count > 15 and len(resolution_text or "") < 50:
                new_flags.append({
                    "id": str(uuid.uuid4()), "type": "bulk_resolve",
                    "detail": f"{today_count} resolutions today with avg text <50 chars",
                    "timestamp": now, "severity": "high"})
            # Generic text check
            if resolution_text and len(resolution_text.strip()) < 30:
                new_flags.append({
                    "id": str(uuid.uuid4()), "type": "generic_text",
                    "detail": f"Resolution text only {len(resolution_text.strip())} chars",
                    "timestamp": now, "severity": "medium"})
            # Cherry-picking check: get officer's resolved priority distribution
            resolved_by_officer = await loop.run_in_executor(executor, lambda: list(
                db.grievances.find({"assigned_officer": officer_name, "status": "resolved",
                                    "created_at": {"$gte": now - timedelta(days=30)}})))
            pri_dist = {}
            for g in resolved_by_officer:
                pri_dist[g.get("priority", "medium")] = pri_dist.get(g.get("priority", "medium"), 0) + 1
            urgent_high = pri_dist.get("urgent", 0) + pri_dist.get("high", 0)
            total_resolved = sum(pri_dist.values())
            if total_resolved > 10 and urgent_high == 0:
                new_flags.append({
                    "id": str(uuid.uuid4()), "type": "cherry_picking",
                    "detail": f"0 urgent/high cases resolved out of {total_resolved} in 30 days",
                    "timestamp": now, "severity": "medium"})
            existing_flags = record.get("anomaly_flags", [])[-20:]
            existing_flags.extend(new_flags)
            await loop.run_in_executor(executor, lambda: db.officer_analytics.update_one(
                {"_id": officer_id},
                {"$set": {"daily_resolution_counts": daily, "priority_distribution": pri_dist,
                          "anomaly_flags": existing_flags[-30:], "baseline_computed_at": now}},
                upsert=True))
        except Exception as e:
            logger.error("Officer anomaly detection error: %s", e)

# ---------------------------------------------------------------------------
# Citizen Impact Scorer
# ---------------------------------------------------------------------------
class ImpactScorer:
    @staticmethod
    async def compute_score(grievance_doc: dict) -> int:
        """Compute a 0-100 impact score based on multiple factors."""
        score = 0
        dept = grievance_doc.get("department", "general")
        district = grievance_doc.get("district")
        # 1. Survival necessity (0-25)
        score += DEPT_SURVIVAL_SCORES.get(dept, 5)
        # 2. Vulnerability index (0-25)
        if district and district in CENSUS_DATA:
            vi = CENSUS_DATA[district].get("vulnerability_index", 0.5)
            score += int(vi * 25)
        else:
            score += 12  # default mid-range
        # 3. Duration/urgency (0-20) — heuristic from description
        desc = grievance_doc.get("description", "").lower()
        duration_score = 10  # default
        if any(w in desc for w in ["months", "year", "long time", "several weeks"]):
            duration_score = 18
        elif any(w in desc for w in ["weeks", "fortnight", "15 days"]):
            duration_score = 14
        elif any(w in desc for w in ["days", "last week", "few days"]):
            duration_score = 8
        elif any(w in desc for w in ["today", "yesterday", "just now"]):
            duration_score = 4
        score += duration_score
        # 4. Recurrence (0-15)
        try:
            loop = asyncio.get_event_loop()
            location = grievance_doc.get("location")
            recurrence_filter = {"department": dept,
                                 "created_at": {"$gte": datetime.now(timezone.utc) - timedelta(days=365)}}
            if district:
                recurrence_filter["district"] = district
            prev_count = await loop.run_in_executor(executor,
                                                     lambda: db.grievances.count_documents(recurrence_filter))
            score += min(prev_count * 3, 15)
        except Exception:
            pass
        # 5. Seasonal urgency (0-15)
        month = datetime.now(timezone.utc).month
        if dept == "rural_water_supply" and month in (3, 4, 5, 6):
            score += 15
        elif dept == "sanitation" and month in (5, 6):
            score += 10
        elif dept == "infrastructure" and month in (6, 7, 8, 9):
            score += 10
        return min(score, 100)

# ---------------------------------------------------------------------------
# Systemic Issue Detection (Constellation Engine)
# ---------------------------------------------------------------------------
class ConstellationEngine:
    @staticmethod
    async def analyze_new_grievance(grievance_doc: dict):
        """Background task: check if this grievance forms part of a systemic cluster."""
        try:
            location = grievance_doc.get("location")
            dept = grievance_doc.get("department")
            if not location or not dept:
                return
            coords = location.get("coordinates", [])
            if len(coords) != 2:
                return
            loop = asyncio.get_event_loop()
            cutoff = datetime.now(timezone.utc) - timedelta(days=30)
            # Find nearby same-department grievances
            nearby = await loop.run_in_executor(executor, lambda: list(db.grievances.find({
                "department": dept,
                "status": {"$in": ["pending", "in_progress", "escalated"]},
                "created_at": {"$gte": cutoff},
                "location": {
                    "$near": {
                        "$geometry": {"type": "Point", "coordinates": coords},
                        "$maxDistance": 10000  # 10km
                    }
                },
                "_id": {"$ne": grievance_doc["_id"]}
            }).limit(20)))
            if len(nearby) < 4:  # need 5 total including current
                return
            # Check semantic similarity
            if not qdrant:
                return
            new_text = f"{grievance_doc['title']} {grievance_doc['description']}"
            new_emb = await get_openai_embedding(new_text)
            similar_ids = [grievance_doc["_id"]]
            for g in nearby:
                if g.get("systemic_issue_id"):
                    continue
                old_text = f"{g['title']} {g['description']}"
                old_emb = await get_openai_embedding(old_text)
                sim = sum(a * b for a, b in zip(new_emb, old_emb)) / (
                    (sum(a**2 for a in new_emb)**0.5) * (sum(b**2 for b in old_emb)**0.5) + 1e-10)
                if sim > 0.75:
                    similar_ids.append(g["_id"])
            if len(similar_ids) < 5:
                return
            # Cluster found — generate root cause analysis via GPT
            cluster_grievances = [grievance_doc] + [g for g in nearby if g["_id"] in similar_ids]
            summaries = [f"- {g['title']}: {g['description'][:200]}" for g in cluster_grievances[:10]]
            prompt = (
                "Analyze these clustered grievances from the same area and department. "
                "They appear to share a common root cause.\n\n"
                + "\n".join(summaries) +
                "\n\nRespond in JSON: {\"title\": \"brief systemic issue title\", "
                "\"root_cause_analysis\": \"2-3 sentence analysis of the probable root cause\", "
                "\"estimated_population_affected\": number}"
            )
            result = await openai_chat([{"role": "user", "content": prompt}], json_mode=True)
            analysis = json.loads(result) if result else {}
            # Compute centroid
            all_coords = [g.get("location", {}).get("coordinates", [0, 0]) for g in cluster_grievances if g.get("location")]
            if all_coords:
                avg_lon = sum(c[0] for c in all_coords) / len(all_coords)
                avg_lat = sum(c[1] for c in all_coords) / len(all_coords)
            else:
                avg_lon, avg_lat = coords[0], coords[1]
            # Get highest priority in cluster
            pri_order = {"urgent": 0, "high": 1, "medium": 2, "low": 3}
            highest_pri = min(cluster_grievances, key=lambda g: pri_order.get(g.get("priority", "medium"), 2))
            issue_doc = {
                "_id": str(uuid.uuid4()),
                "title": analysis.get("title", f"Systemic {dept} issue in {grievance_doc.get('district', 'unknown')}"),
                "root_cause_analysis": analysis.get("root_cause_analysis", "Multiple related complaints detected in close proximity."),
                "department": dept,
                "district": grievance_doc.get("district"),
                "affected_area_center": {"type": "Point", "coordinates": [avg_lon, avg_lat]},
                "affected_radius_km": 10.0,
                "grievance_ids": similar_ids,
                "estimated_population_affected": analysis.get("estimated_population_affected", 0),
                "priority": highest_pri.get("priority", "high"),
                "status": "detected",
                "created_at": datetime.now(timezone.utc),
                "updated_at": datetime.now(timezone.utc),
                "assigned_officer": None,
            }
            await loop.run_in_executor(executor, db.systemic_issues.insert_one, issue_doc)
            # Mark all grievances with systemic_issue_id
            await loop.run_in_executor(executor, lambda: db.grievances.update_many(
                {"_id": {"$in": similar_ids}},
                {"$set": {"systemic_issue_id": issue_doc["_id"]}}))
            logger.info("Systemic issue detected: %s with %d grievances", issue_doc["title"], len(similar_ids))
        except Exception as e:
            logger.error("Constellation engine error: %s", e)

# ---------------------------------------------------------------------------
# Predictive Forecaster
# ---------------------------------------------------------------------------
class PredictiveForecaster:
    @staticmethod
    async def generate_forecast(weeks_ahead: int = 2) -> dict:
        """Generate a predictive forecast based on historical data and external signals."""
        loop = asyncio.get_event_loop()
        now = datetime.now(timezone.utc)
        # Step 1: Historical baselines (past 1 year by district+department+week)
        year_ago = now - timedelta(days=365)
        all_grievances = await loop.run_in_executor(executor, lambda: list(
            db.grievances.find({"created_at": {"$gte": year_ago}},
                               {"district": 1, "department": 1, "created_at": 1}).limit(10000)))
        baselines = {}
        for g in all_grievances:
            d = g.get("district") or "Unknown"
            dept = g.get("department") or "general"
            week = g["created_at"].isocalendar()[1]
            key = (d, dept, week)
            baselines[key] = baselines.get(key, 0) + 1
        # Step 2: Seasonal baseline for the target week
        target_week = (now + timedelta(weeks=weeks_ahead)).isocalendar()[1]
        predictions = []
        districts_seen = set()
        for (d, dept, week), count in baselines.items():
            districts_seen.add(d)
            if week == target_week:
                predictions.append({
                    "district": d, "department": dept,
                    "baseline_count": count, "predicted_count": count,
                    "spike_factor": 1.0, "confidence": 0.5,
                    "triggers": ["seasonal_baseline"],
                    "recommended_actions": []
                })
        # Step 3: Weather correlation (if API key available)
        weather_data = {}
        if OPENWEATHERMAP_API_KEY:
            async with httpx.AsyncClient(timeout=10.0) as client:
                for district in list(districts_seen)[:10]:
                    try:
                        resp = await client.get(
                            f"https://api.openweathermap.org/data/2.5/forecast",
                            params={"q": f"{district},IN", "appid": OPENWEATHERMAP_API_KEY, "units": "metric"})
                        if resp.status_code == 200:
                            weather_data[district] = resp.json()
                    except Exception:
                        pass
        # Step 4: Apply weather multipliers
        for pred in predictions:
            d = pred["district"]
            if d in weather_data:
                forecasts_list = weather_data[d].get("list", [])
                max_rain = max((f.get("rain", {}).get("3h", 0) for f in forecasts_list), default=0)
                max_temp = max((f.get("main", {}).get("temp_max", 30) for f in forecasts_list), default=30)
                if max_rain > 50:
                    if pred["department"] in ("rural_water_supply", "infrastructure"):
                        pred["predicted_count"] = int(pred["baseline_count"] * 2.5)
                        pred["spike_factor"] = 2.5
                        pred["triggers"].append("heavy_rainfall")
                        pred["confidence"] = 0.75
                elif max_rain > 20:
                    if pred["department"] in ("rural_water_supply", "infrastructure", "sanitation"):
                        pred["predicted_count"] = int(pred["baseline_count"] * 1.5)
                        pred["spike_factor"] = 1.5
                        pred["triggers"].append("moderate_rainfall")
                if max_temp > 42:
                    if pred["department"] == "rural_water_supply":
                        pred["predicted_count"] = int(pred["predicted_count"] * 2)
                        pred["spike_factor"] *= 2
                        pred["triggers"].append("heat_wave")
                        pred["confidence"] = 0.8
        # Step 5: Seasonal calendar rules
        month = (now + timedelta(weeks=weeks_ahead)).month
        for pred in predictions:
            if month in (5, 6) and pred["department"] == "sanitation":
                pred["predicted_count"] = int(pred["predicted_count"] * 1.5)
                pred["spike_factor"] *= 1.5
                pred["triggers"].append("pre_monsoon_season")
        # Step 6: Generate recommended actions via GPT for top spikes
        predictions.sort(key=lambda p: p["spike_factor"], reverse=True)
        top_preds = [p for p in predictions if p["spike_factor"] > 1.2][:5]
        if top_preds:
            pred_summary = "\n".join([
                f"- {p['district']}, {p['department']}: {p['spike_factor']:.1f}x spike expected, triggers: {', '.join(p['triggers'])}"
                for p in top_preds])
            prompt = (
                "You are an advisor for the Odisha Panchayati Raj & Drinking Water Department. "
                "Based on these predicted grievance spikes for the next 2 weeks, suggest 1-2 specific "
                "recommended actions for each prediction.\n\n" + pred_summary +
                "\n\nRespond in JSON: {\"recommendations\": [{\"district\": \"...\", \"department\": \"...\", \"actions\": [\"...\", \"...\"]}]}"
            )
            try:
                result = await openai_chat([{"role": "user", "content": prompt}], json_mode=True)
                if result:
                    recs = json.loads(result).get("recommendations", [])
                    rec_map = {(r["district"], r["department"]): r.get("actions", []) for r in recs}
                    for p in top_preds:
                        p["recommended_actions"] = rec_map.get((p["district"], p["department"]), [])
            except Exception:
                pass
        forecast_doc = {
            "_id": str(uuid.uuid4()),
            "forecast_date": now,
            "forecast_period_start": now + timedelta(weeks=1),
            "forecast_period_end": now + timedelta(weeks=1 + weeks_ahead),
            "predictions": predictions[:50],
            "weather_data_used": {d: True for d in weather_data},
            "generated_by": "system",
        }
        await loop.run_in_executor(executor, db.forecasts.insert_one, forecast_doc)
        return forecast_doc

# ---------------------------------------------------------------------------
# Scheme-Grievance Matcher
# ---------------------------------------------------------------------------
class SchemeGrievanceMatcher:
    @staticmethod
    async def find_matching_scheme(title: str, description: str, department: str) -> Optional[Dict]:
        """Check if any scheme in the knowledge base matches this grievance."""
        if not qdrant:
            return None
        try:
            text = f"{title} {description}"
            embedding = await get_openai_embedding(text)
            dept_filter = Filter(must=[FieldCondition(key="department", match=MatchValue(value=department))])
            results = qdrant.query_points(
                collection_name="schemes", query=embedding, query_filter=dept_filter,
                limit=1, with_payload=True, with_vectors=False)
            if not results.points:
                return None
            top = results.points[0]
            if top.score < 0.75:
                return None
            scheme_name = top.payload.get("name", "Unknown Scheme")
            eligibility = top.payload.get("eligibility", "")
            # Quick eligibility check via GPT
            prompt = (
                f"Given this grievance: \"{title} - {description[:500]}\"\n"
                f"And this government scheme: {scheme_name}\nEligibility: {eligibility}\n\n"
                "Is the citizen likely eligible for this scheme? "
                "Respond in JSON: {\"eligible\": true/false, \"reasoning\": \"one sentence\"}"
            )
            result = await openai_chat([{"role": "user", "content": prompt}], json_mode=True)
            elig_data = json.loads(result) if result else {"eligible": False, "reasoning": ""}
            return {
                "scheme_id": str(top.id), "scheme_name": scheme_name,
                "relevance_score": round(top.score, 3),
                "eligibility_likely": elig_data.get("eligible", False),
                "eligibility_reasoning": elig_data.get("reasoning", ""),
            }
        except Exception as e:
            logger.error("Scheme matching error: %s", e)
            return None

# ---------------------------------------------------------------------------
# Cross-Department Dependency Analyzer
# ---------------------------------------------------------------------------
class CrossDeptAnalyzer:
    @staticmethod
    async def analyze_dependencies(title: str, description: str) -> Optional[List[Dict]]:
        """Determine if a grievance requires multi-department coordination."""
        try:
            prompt = (
                "You are an analyst for the Odisha Panchayati Raj & Drinking Water Department.\n"
                "Departments: panchayati_raj, rural_water_supply, mgnregs, rural_housing, "
                "rural_livelihoods, sanitation, infrastructure, general\n\n"
                f"Grievance: {title} - {description[:1000]}\n\n"
                "Does this grievance require action from MULTIPLE departments? "
                "Only say yes if it genuinely needs coordination (e.g., waterlogged road needs drainage + road repair).\n"
                "Respond in JSON:\n"
                "{\"is_multi_department\": true/false, \"sub_tasks\": ["
                "{\"department\": \"...\", \"task\": \"...\", \"depends_on\": [\"dept_name_or_empty\"]}]}"
            )
            result = await openai_chat([{"role": "user", "content": prompt}], json_mode=True)
            if not result:
                return None
            data = json.loads(result)
            if not data.get("is_multi_department"):
                return None
            sub_tasks = []
            for st in data.get("sub_tasks", []):
                sub_tasks.append({
                    "id": str(uuid.uuid4()),
                    "department": st.get("department", "general"),
                    "task": st.get("task", ""),
                    "depends_on": st.get("depends_on", []),
                    "status": "pending",
                    "assigned_officer": None,
                })
            return sub_tasks if len(sub_tasks) >= 2 else None
        except Exception as e:
            logger.error("Cross-dept analysis error: %s", e)
            return None

# ---------------------------------------------------------------------------
# Service Memory Deadline Estimator
# ---------------------------------------------------------------------------
class DeadlineEstimator:
    @staticmethod
    async def estimate_from_service_memory(title: str, description: str, created_at: datetime) -> Optional[datetime]:
        """Query service memory for similar resolved cases and estimate resolution deadline."""
        if not qdrant:
            return None
        try:
            text = f"{title} {description}"
            embedding = await get_openai_embedding(text)
            results = qdrant.query_points(
                collection_name="service_memory", query=embedding,
                limit=5, with_payload=True, with_vectors=False)
            if not results.points:
                return None
            durations = []
            loop = asyncio.get_event_loop()
            for p in results.points:
                gid = p.payload.get("grievance_id")
                if not gid:
                    continue
                g = await loop.run_in_executor(executor, db.grievances.find_one,
                                               {"_id": gid, "status": "resolved"})
                if g and g.get("updated_at") and g.get("created_at"):
                    dur = (g["updated_at"] - g["created_at"]).total_seconds()
                    if dur > 0:
                        durations.append(dur)
            if not durations:
                return None
            median_secs = statistics.median(durations)
            return created_at + timedelta(seconds=median_secs)
        except Exception as e:
            logger.error("Deadline estimation error: %s", e)
            return None

# ---------------------------------------------------------------------------
# Core Grievance Processing
# ---------------------------------------------------------------------------
SELF_RESOLVE_CONFIDENCE_THRESHOLD = 0.60

def find_officer_for_department(db, department: str) -> Optional[str]:
    """Find an officer assigned to a department and return their full_name, or None."""
    if not department:
        return None
    officer = db.users.find_one({"role": "officer", "department": department})
    return officer["full_name"] if officer else None

async def process_grievance(data: GrievanceCreate, db, user: Optional[dict] = None,
                            attachment_ids: Optional[List[str]] = None) -> GrievanceResponse:
    text = f"{data.title} {data.description}"
    sentiment = await SentimentAnalyzer.analyze_sentiment(text)
    department, priority, officer_category, tier = await GrievanceClassifier.classify(data.title, data.description)
    if sentiment == Sentiment.FRUSTRATED and priority in [Priority.LOW, Priority.MEDIUM]:
        priority = Priority.HIGH

    # Override tier to escalation for urgent+frustrated
    if sentiment == Sentiment.FRUSTRATED and priority == Priority.URGENT:
        tier = ResolutionTier.ESCALATION
    if tier == ResolutionTier.ESCALATION:
        priority = max(priority, Priority.HIGH, key=lambda p: [Priority.LOW, Priority.MEDIUM, Priority.HIGH, Priority.URGENT].index(p))

    doc_results = await KnowledgeSearcher.search_documentation(text)
    memory_results = await KnowledgeSearcher.search_service_memory(text)
    scheme_results = await KnowledgeSearcher.search_schemes(text)
    all_context = memory_results + doc_results + scheme_results

    ai_resolution = None
    confidence_score = 0.0
    citizen_name = data.citizen_name if not data.is_anonymous else None

    if all_context:
        if tier == ResolutionTier.SELF_RESOLVABLE:
            ai_resolution, confidence_score = await ResolutionGenerator.generate_citizen_guidance(
                data.title, data.description, citizen_name,
                sentiment.value, data.language.value, all_context)
            if confidence_score < SELF_RESOLVE_CONFIDENCE_THRESHOLD:
                tier = ResolutionTier.OFFICER_ACTION
        else:
            ai_resolution, confidence_score = await ResolutionGenerator.generate_officer_draft(
                data.title, data.description, citizen_name,
                sentiment.value, data.language.value, all_context)

    g_status = GrievanceStatus.ESCALATED if tier == ResolutionTier.ESCALATION else GrievanceStatus.PENDING

    # Auto-assign officer when escalated at creation
    assigned_officer_name = None
    if g_status == GrievanceStatus.ESCALATED:
        assigned_officer_name = find_officer_for_department(db, department.value)

    tracking = generate_tracking_number(db)
    sla = calculate_sla_deadline(priority)
    now = datetime.now(timezone.utc)

    # Estimate resolution deadline from service memory
    estimated_deadline = await DeadlineEstimator.estimate_from_service_memory(data.title, data.description, now)

    # Scheme-grievance matching
    scheme_match = await SchemeGrievanceMatcher.find_matching_scheme(data.title, data.description, department.value)

    # Cross-department dependency analysis
    sub_tasks = await CrossDeptAnalyzer.analyze_dependencies(data.title, data.description)

    doc = {
        "_id": str(uuid.uuid4()), "tracking_number": tracking,
        "title": data.title, "description": data.description,
        "citizen_name": None if data.is_anonymous else data.citizen_name,
        "citizen_email": None if data.is_anonymous else data.citizen_email,
        "citizen_phone": None if data.is_anonymous else data.citizen_phone,
        "is_anonymous": data.is_anonymous,
        "is_public": data.is_public and not data.is_anonymous,
        "language": data.language.value,
        "district": data.district, "department": department.value,
        "priority": priority.value, "officer_category": officer_category.value,
        "status": g_status.value, "sentiment": sentiment.value,
        "resolution_tier": tier.value,
        "created_at": now, "updated_at": now,
        "sla_deadline": sla,
        "estimated_resolution_deadline": estimated_deadline,
        "ai_resolution": ai_resolution,
        "manual_resolution": None, "resolution_type": None,
        "confidence_score": confidence_score, "assigned_officer": assigned_officer_name,
        "resolution_feedback": None, "notes": [],
        "location": {"type": "Point", "coordinates": [data.location.longitude, data.location.latitude]}
            if data.location and data.location.longitude and data.location.latitude else None,
        "citizen_user_id": str(user["_id"]) if user else None,
        "attachments": attachment_ids or [],
        "scheme_match": scheme_match,
        "sub_tasks": sub_tasks or [],
        "systemic_issue_id": None,
        "impact_score": None,
    }
    loop = asyncio.get_event_loop()
    await loop.run_in_executor(executor, db.grievances.insert_one, doc)

    # Compute impact score (needs doc in DB for recurrence check)
    impact_score = await ImpactScorer.compute_score(doc)
    await loop.run_in_executor(executor, lambda: db.grievances.update_one(
        {"_id": doc["_id"]}, {"$set": {"impact_score": impact_score}}))
    doc["impact_score"] = impact_score

    return GrievanceResponse(**doc, id=doc["_id"])

# ---------------------------------------------------------------------------
# AUTH ENDPOINTS
# ---------------------------------------------------------------------------
@app.post("/auth/register", response_model=TokenResponse)
@limiter.limit("3/minute")
async def register(request: Request, user_data: UserCreate, db=Depends(get_db)):
    # Public registration is citizen-only; officers/admins created via admin panel
    if user_data.role != UserRole.CITIZEN:
        raise HTTPException(status_code=403, detail="Public registration is for citizens only. Officer/admin accounts must be created by an administrator.")
    loop = asyncio.get_event_loop()
    existing = await loop.run_in_executor(executor, db.users.find_one, {"username": user_data.username})
    if existing:
        raise HTTPException(status_code=400, detail="Username already exists")
    user_doc = {
        "_id": str(uuid.uuid4()), "username": user_data.username,
        "hashed_password": hash_password(user_data.password),
        "full_name": user_data.full_name, "email": user_data.email,
        "phone": user_data.phone, "role": user_data.role.value,
        "department": user_data.department.value if user_data.department else None,
        "created_at": datetime.now(timezone.utc),
        "face_descriptor": user_data.face_descriptor,
    }
    
    # Check if face is already registered
    if user_data.face_descriptor:
        existing_face_user = await find_user_by_face(user_data.face_descriptor, db)
        if existing_face_user:
            raise HTTPException(status_code=400, detail="This face is already registered with another account.")
            
    await loop.run_in_executor(executor, db.users.insert_one, user_doc)
    token = create_access_token({"sub": user_data.username, "role": user_data.role.value})
    return TokenResponse(access_token=token, user=user_to_response(user_doc))

@app.post("/auth/login", response_model=TokenResponse)
@limiter.limit("30/minute")
async def login(request: Request, form: UserLogin, db=Depends(get_db)):
    loop = asyncio.get_event_loop()

    # Face Login
    # Face Login
    if form.face_descriptor:
        best_match = await find_user_by_face(form.face_descriptor, db)
        if best_match:
            token = create_access_token({"sub": best_match["username"], "role": best_match["role"]})
            return TokenResponse(access_token=token, user=user_to_response(best_match))
        
        raise HTTPException(status_code=401, detail="Face not recognized or match confidence too low")

    # Password Login
    if not form.username or not form.password:
        raise HTTPException(status_code=400, detail="Username and password required")

    user = await loop.run_in_executor(executor, db.users.find_one, {"username": form.username})
    if not user or not verify_password(form.password, user["hashed_password"]):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    token = create_access_token({"sub": user["username"], "role": user["role"]})
    return TokenResponse(access_token=token, user=user_to_response(user))

@app.get("/auth/me", response_model=UserResponse)
async def get_me(user=Depends(get_current_user)):
    return user_to_response(user)

@app.middleware("http")
async def add_permissions_policy(request, call_next):
    response = await call_next(request)
    response.headers["Permissions-Policy"] = "camera=*, microphone=*"
    response.headers["Feature-Policy"] = "camera *; microphone *"
    return response
    
@app.post("/auth/logout")
async def logout(token: Optional[str] = Depends(oauth2_scheme)):
    if token:
        _token_blacklist.add(token)
        # Prune blacklist if it grows too large (expired tokens don't matter)
        if len(_token_blacklist) > 10000:
            _token_blacklist.clear()
    return {"detail": "Logged out successfully"}

@app.get("/auth/officers", response_model=List[UserResponse])
async def list_officers(user=Depends(require_role(UserRole.OFFICER.value, UserRole.ADMIN.value)),
                        db=Depends(get_db)):
    loop = asyncio.get_event_loop()
    officers = await loop.run_in_executor(
        executor, lambda: list(db.users.find({"role": {"$in": ["officer", "admin"]}})))
    return [user_to_response(o) for o in officers]

# ---------------------------------------------------------------------------
# ADMIN USER MANAGEMENT ENDPOINTS
# ---------------------------------------------------------------------------
@app.get("/admin/users", response_model=List[UserResponse])
async def admin_list_users(role: Optional[str] = None,
                           user=Depends(require_role(UserRole.ADMIN.value)),
                           db=Depends(get_db)):
    query = {}
    if role and role in [r.value for r in UserRole]:
        query["role"] = role
    loop = asyncio.get_event_loop()
    users = await loop.run_in_executor(executor, lambda: list(db.users.find(query).sort("created_at", -1)))
    return [user_to_response(u) for u in users]

@app.post("/admin/users", response_model=UserResponse)
async def admin_create_user(user_data: UserCreate,
                            user=Depends(require_role(UserRole.ADMIN.value)),
                            db=Depends(get_db)):
    loop = asyncio.get_event_loop()
    existing = await loop.run_in_executor(executor, db.users.find_one, {"username": user_data.username})
    if existing:
        raise HTTPException(status_code=400, detail="Username already exists")
    user_doc = {
        "_id": str(uuid.uuid4()), "username": user_data.username,
        "hashed_password": hash_password(user_data.password),
        "full_name": user_data.full_name, "email": user_data.email,
        "phone": user_data.phone, "role": user_data.role.value,
        "department": user_data.department.value if user_data.department else None,
        "created_at": datetime.now(timezone.utc),
        "face_descriptor": user_data.face_descriptor,
    }

    # Check if face is already registered
    if user_data.face_descriptor:
        existing_face_user = await find_user_by_face(user_data.face_descriptor, db)
        if existing_face_user:
            raise HTTPException(status_code=400, detail="This face is already registered with another account.")

    await loop.run_in_executor(executor, db.users.insert_one, user_doc)
    logger.info("Admin %s created user %s (%s)", user["username"], user_data.username, user_data.role.value)
    return user_to_response(user_doc)

@app.put("/admin/users/{user_id}/block")
async def admin_block_user(user_id: str, action: str = Query(...),
                           reason: str = Query("Manual admin action"),
                           user=Depends(require_role(UserRole.ADMIN.value)), db=Depends(get_db)):
    """Admin can manually block, unblock, or flag a citizen as spam."""
    if action not in ("block", "unblock", "flag_spam"):
        raise HTTPException(status_code=400, detail="action must be block, unblock, or flag_spam")
    user_id = validate_uuid(user_id, "user_id")
    loop = asyncio.get_event_loop()
    target = await loop.run_in_executor(executor, db.users.find_one, {"_id": user_id})
    if not target:
        raise HTTPException(status_code=404, detail="User not found")
    if target["role"] != UserRole.CITIZEN.value:
        raise HTTPException(status_code=400, detail="Can only block/flag citizen accounts")

    now = datetime.now(timezone.utc)
    if action == "block":
        await loop.run_in_executor(executor, lambda: db.spam_tracking.update_one(
            {"_id": user_id},
            {"$set": {"is_blocked": True, "blocked_at": now, "photo_id_status": "none"},
             "$push": {"pattern_flags": {"type": "manual_block", "detail": reason,
                                          "by": user["username"], "at": now.isoformat()}}},
            upsert=True))
        logger.info("Admin %s blocked user %s: %s", user["username"], target["username"], reason)
        return {"detail": f"User '{target['username']}' blocked"}
    elif action == "unblock":
        await loop.run_in_executor(executor, lambda: db.spam_tracking.update_one(
            {"_id": user_id},
            {"$set": {"is_blocked": False, "spam_score": 0.0, "photo_id_status": "approved",
                      "pattern_flags": []}},
            upsert=True))
        logger.info("Admin %s unblocked user %s", user["username"], target["username"])
        return {"detail": f"User '{target['username']}' unblocked"}
    else:  # flag_spam
        await loop.run_in_executor(executor, lambda: db.spam_tracking.update_one(
            {"_id": user_id},
            {"$set": {"is_blocked": True, "blocked_at": now, "photo_id_status": "none"},
             "$push": {"pattern_flags": {"type": "admin_flag_spam", "detail": reason,
                                          "by": user["username"], "at": now.isoformat()}}},
            upsert=True))
        logger.info("Admin %s flagged user %s as spam: %s", user["username"], target["username"], reason)
        return {"detail": f"User '{target['username']}' flagged — must upload photo ID"}

@app.put("/admin/users/{user_id}", response_model=UserResponse)
async def admin_update_user(user_id: str, update: UserUpdate,
                            user=Depends(require_role(UserRole.ADMIN.value)),
                            db=Depends(get_db)):
    user_id = validate_uuid(user_id, "user_id")
    loop = asyncio.get_event_loop()
    target = await loop.run_in_executor(executor, db.users.find_one, {"_id": user_id})
    if not target:
        raise HTTPException(status_code=404, detail="User not found")
    set_fields: Dict[str, Any] = {}
    if update.full_name is not None:
        set_fields["full_name"] = update.full_name
    if update.email is not None:
        set_fields["email"] = update.email
    if update.phone is not None:
        set_fields["phone"] = update.phone
    if update.role is not None:
        set_fields["role"] = update.role.value
    if update.department is not None:
        set_fields["department"] = update.department.value
    if update.password is not None:
        set_fields["hashed_password"] = hash_password(update.password)
    
    if update.face_descriptor is not None:
        # If setting a new face, check uniqueness
        if update.face_descriptor: # if not empty list or None
             existing_face_user = await find_user_by_face(update.face_descriptor, db)
             if existing_face_user and str(existing_face_user["_id"]) != user_id:
                 raise HTTPException(status_code=400, detail="This face is already registered with another account.")
        set_fields["face_descriptor"] = update.face_descriptor

    if not set_fields:
        raise HTTPException(status_code=400, detail="No fields to update")
    result = await loop.run_in_executor(
        executor, lambda: db.users.update_one({"_id": user_id}, {"$set": set_fields}))
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="User not found")
    updated = await loop.run_in_executor(executor, db.users.find_one, {"_id": user_id})
    logger.info("Admin %s updated user %s", user["username"], target["username"])
    return user_to_response(updated)

@app.delete("/admin/users/{user_id}")
async def admin_delete_user(user_id: str,
                            user=Depends(require_role(UserRole.ADMIN.value)),
                            db=Depends(get_db)):
    user_id = validate_uuid(user_id, "user_id")
    if str(user["_id"]) == user_id:
        raise HTTPException(status_code=400, detail="Cannot delete your own account")
    loop = asyncio.get_event_loop()
    target = await loop.run_in_executor(executor, db.users.find_one, {"_id": user_id})
    if not target:
        raise HTTPException(status_code=404, detail="User not found")
    await loop.run_in_executor(executor, db.users.delete_one, {"_id": user_id})
    logger.info("Admin %s deleted user %s", user["username"], target["username"])
    return {"detail": f"User '{target['username']}' deleted"}

# ---------------------------------------------------------------------------
# ADMIN DEADLINE ALERTS ENDPOINT
# ---------------------------------------------------------------------------
@app.get("/admin/deadline-alerts")
async def admin_deadline_alerts(
    user=Depends(require_role(UserRole.ADMIN.value)),
    db=Depends(get_db),
):
    """Return unresolved grievances nearing or past their SLA deadline."""
    now = datetime.now(timezone.utc)
    threshold_24h = now + timedelta(hours=24)
    threshold_48h = now + timedelta(hours=48)

    def fetch():
        # All unresolved grievances that have an SLA deadline within 48h or already breached
        cursor = db.grievances.find(
            {
                "status": {"$in": ["pending", "in_progress", "escalated"]},
                "sla_deadline": {"$exists": True, "$ne": None, "$lte": threshold_48h},
            },
            {
                "_id": 1, "tracking_number": 1, "title": 1, "department": 1,
                "priority": 1, "status": 1, "sla_deadline": 1, "created_at": 1,
            },
        ).sort("sla_deadline", 1)
        return list(cursor)

    loop = asyncio.get_event_loop()
    results = await loop.run_in_executor(executor, fetch)

    breached, critical, warning = [], [], []
    for g in results:
        item = {
            "id": g["_id"],
            "tracking_number": g["tracking_number"],
            "title": g["title"],
            "department": g.get("department", "general"),
            "priority": g.get("priority", "medium"),
            "status": g["status"],
            "sla_deadline": g["sla_deadline"].isoformat() if g.get("sla_deadline") else None,
            "created_at": g["created_at"].isoformat() if g.get("created_at") else None,
        }
        deadline = g["sla_deadline"]
        # Ensure both sides are tz-aware for comparison (MongoDB may store naive datetimes)
        if deadline.tzinfo is None:
            deadline = deadline.replace(tzinfo=timezone.utc)
        if deadline <= now:
            breached.append(item)
        elif deadline <= threshold_24h:
            critical.append(item)
        else:
            warning.append(item)

    return {
        "breached": breached,
        "critical": critical,
        "warning": warning,
        "total_alerts": len(breached) + len(critical) + len(warning),
    }

# ---------------------------------------------------------------------------
# GRIEVANCE ENDPOINTS
# ---------------------------------------------------------------------------
@app.post("/grievances", response_model=GrievanceResponse)
async def create_grievance(
    request: Request,
    background_tasks: BackgroundTasks,
    user=Depends(get_optional_user),
    db=Depends(get_db),
    # Form fields (used when submitting with files via multipart/form-data)
    title: Optional[str] = Form(None),
    description: Optional[str] = Form(None),
    language: Optional[str] = Form(None),
    district: Optional[str] = Form(None),
    is_anonymous: Optional[str] = Form(None),
    is_public: Optional[str] = Form(None),
    citizen_name: Optional[str] = Form(None),
    citizen_email: Optional[str] = Form(None),
    citizen_phone: Optional[str] = Form(None),
    latitude: Optional[str] = Form(None),
    longitude: Optional[str] = Form(None),
    files: List[UploadFile] = File(default=[]),
):
    try:
        client_ip = request.client.host if request.client else ""
        # Determine if this is a form submission or JSON body
        ct = request.headers.get("content-type", "")
        if title is not None and description is not None:
            # multipart/form-data submission
            loc = None
            try:
                lat_f = float(latitude) if latitude else None
                lon_f = float(longitude) if longitude else None
                if lat_f is not None and lon_f is not None:
                    loc = Geolocation(latitude=lat_f, longitude=lon_f)
            except (ValueError, TypeError):
                pass
            data = GrievanceCreate(
                title=title, description=description,
                language=Language(language) if language else Language.ENGLISH,
                district=district or None,
                is_anonymous=is_anonymous in ("true", "True", "1", "on") if is_anonymous else False,
                is_public=is_public in ("true", "True", "1", "on") if is_public else False,
                citizen_name=citizen_name, citizen_email=citizen_email,
                citizen_phone=citizen_phone, location=loc,
            )
        else:
            # JSON body fallback
            body = await request.json()
            data = GrievanceCreate(**body)

        # Spam check
        if user:
            is_blocked, reason = await SpamDetector.check_and_update(
                str(user["_id"]), data.title, data.description, client_ip)
            if is_blocked:
                raise HTTPException(status_code=403, detail=reason)

        # Handle file uploads
        attachment_ids = []
        if files:
            if len(files) > MAX_FILES_PER_GRIEVANCE:
                raise HTTPException(status_code=400,
                                    detail=f"Maximum {MAX_FILES_PER_GRIEVANCE} files allowed")
            loop = asyncio.get_event_loop()
            for f in files:
                if f.content_type and f.content_type not in ALLOWED_MIME_TYPES:
                    raise HTTPException(status_code=400,
                                        detail=f"File type {f.content_type} not allowed")
                content = await f.read()
                if len(content) > MAX_FILE_SIZE:
                    raise HTTPException(status_code=400,
                                        detail=f"File {f.filename} exceeds 50MB limit")
                fid = await loop.run_in_executor(executor, lambda c=content, fn=f.filename, ct=f.content_type: gfs.put(
                    c, filename=fn, content_type=ct,
                    metadata={"uploaded_at": datetime.now(timezone.utc).isoformat()}))
                attachment_ids.append(str(fid))

        result = await process_grievance(data, db, user, attachment_ids=attachment_ids)

        # Background tasks: pattern analysis and constellation engine
        if user:
            background_tasks.add_task(PatternAnalyzer.analyze,
                                      str(user["_id"]), data.title, data.description, client_ip)
        # Get the full doc for constellation engine
        loop = asyncio.get_event_loop()
        full_doc = await loop.run_in_executor(executor, db.grievances.find_one, {"_id": result.id})
        if full_doc:
            background_tasks.add_task(ConstellationEngine.analyze_new_grievance, full_doc)

        # Strip officer-only AI drafts for citizens
        is_citizen = user is None or user.get("role") == UserRole.CITIZEN.value
        if is_citizen and result.resolution_tier != ResolutionTier.SELF_RESOLVABLE:
            result.ai_resolution = None
            result.confidence_score = None
        return result
    except HTTPException:
        raise
    except (ValueError, TypeError) as e:
        # Pydantic validation / type coercion errors → 422
        logger.warning("Grievance validation error: %s", e)
        raise HTTPException(status_code=422, detail=str(e))
    except Exception as e:
        logger.error("Error creating grievance: %s", e)
        raise HTTPException(status_code=500, detail="Internal server error")

@app.get("/grievances", response_model=List[GrievanceResponse])
async def get_grievances(
    status: Optional[GrievanceStatus] = None, department: Optional[Department] = None,
    priority: Optional[Priority] = None, district: Optional[str] = None,
    group_by_date: bool = Query(False),
    sort_by: Optional[str] = Query(None),
    limit: int = Query(50, ge=1, le=100), skip: int = Query(0, ge=0, le=10000),
    user=Depends(get_current_user), db=Depends(get_db)):
    try:
        fq = {}
        if status: fq["status"] = status.value
        if department: fq["department"] = department.value
        if priority: fq["priority"] = priority.value
        if district:
            district = sanitize_str(district)
            if district not in ODISHA_DISTRICTS:
                raise HTTPException(status_code=400, detail="Invalid district name")
            fq["district"] = district
        if user["role"] == UserRole.CITIZEN.value:
            fq["citizen_user_id"] = str(user["_id"])

        if group_by_date:
            # Aggregation: group by date, sort by priority within each date
            def fetch_grouped():
                pipeline = [
                    {"$match": fq},
                    {"$addFields": {
                        "date_group": {"$dateToString": {"format": "%Y-%m-%d", "date": "$created_at"}},
                        "priority_order": {"$switch": {
                            "branches": [
                                {"case": {"$eq": ["$priority", "urgent"]}, "then": 0},
                                {"case": {"$eq": ["$priority", "high"]}, "then": 1},
                                {"case": {"$eq": ["$priority", "medium"]}, "then": 2},
                            ], "default": 3
                        }}
                    }},
                    {"$sort": {"date_group": -1, "priority_order": 1}},
                    {"$skip": skip}, {"$limit": limit}
                ]
                return [convert_db_grievance(g) for g in db.grievances.aggregate(pipeline)]
            loop = asyncio.get_event_loop()
            grievances = await loop.run_in_executor(executor, fetch_grouped)
        elif sort_by == "impact_score":
            def fetch_by_impact():
                return [convert_db_grievance(g) for g in
                        db.grievances.find(fq).sort("impact_score", -1).skip(skip).limit(limit)]
            loop = asyncio.get_event_loop()
            grievances = await loop.run_in_executor(executor, fetch_by_impact)
        else:
            def fetch():
                return [convert_db_grievance(g) for g in
                        db.grievances.find(fq).sort("created_at", -1).skip(skip).limit(limit)]
            loop = asyncio.get_event_loop()
            grievances = await loop.run_in_executor(executor, fetch)

        if user["role"] == UserRole.CITIZEN.value:
            for g in grievances:
                if g.get("resolution_tier") != "self_resolvable":
                    g["ai_resolution"] = None
                    g["confidence_score"] = None
        return [GrievanceResponse(**g, id=g["_id"]) for g in grievances]
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Error getting grievances: %s", e)
        raise HTTPException(status_code=500, detail="Internal server error")

@app.get("/grievances/track/{tracking_number}", response_model=GrievanceTrackResponse)
@limiter.limit("10/minute")
async def track_grievance(request: Request, tracking_number: str, db=Depends(get_db)):
    tracking_number = sanitize_str(tracking_number)
    if not re.match(r'^GRV-\d{4}-[A-Za-z0-9]{6,16}$', tracking_number):
        raise HTTPException(status_code=400, detail="Invalid tracking number format")
    loop = asyncio.get_event_loop()
    g = await loop.run_in_executor(executor, db.grievances.find_one, {"tracking_number": tracking_number})
    if not g:
        raise HTTPException(status_code=404, detail="Grievance not found")
    # Only expose AI resolution to the public tracking endpoint for self-resolvable grievances;
    # officer-action / escalation AI drafts are internal and must not be leaked.
    tier = g.get("resolution_tier", "officer_action")
    ai_res = g.get("ai_resolution") if tier == "self_resolvable" else None
    return GrievanceTrackResponse(
        id=g["_id"], tracking_number=g["tracking_number"], title=g["title"], status=g["status"],
        department=g["department"], priority=g["priority"], created_at=g["created_at"],
        updated_at=g["updated_at"], sla_deadline=g.get("sla_deadline"),
        ai_resolution=ai_res, manual_resolution=g.get("manual_resolution"),
        resolution_type=g.get("resolution_type"),
        resolution_tier=g.get("resolution_tier"))

@app.get("/grievances/{grievance_id}", response_model=GrievanceResponse)
async def get_grievance(grievance_id: str, user=Depends(get_current_user), db=Depends(get_db)):
    grievance_id = validate_uuid(grievance_id, "grievance_id")
    loop = asyncio.get_event_loop()
    g = await loop.run_in_executor(executor, db.grievances.find_one, {"_id": grievance_id})
    if not g:
        raise HTTPException(status_code=404, detail="Grievance not found")
    if user["role"] == UserRole.CITIZEN.value and g.get("citizen_user_id") != str(user["_id"]):
        raise HTTPException(status_code=403, detail="Access denied")
    g = convert_db_grievance(g)
    # Strip officer-only AI drafts for citizen users
    if user["role"] == UserRole.CITIZEN.value and g.get("resolution_tier") != "self_resolvable":
        g["ai_resolution"] = None
        g["confidence_score"] = None
    return GrievanceResponse(**g, id=g["_id"])

@app.put("/grievances/{grievance_id}/resolve", response_model=GrievanceResponse)
async def resolve_grievance(grievance_id: str, resolution: ManualResolution,
                             background_tasks: BackgroundTasks,
                             user=Depends(require_role(UserRole.OFFICER.value, UserRole.ADMIN.value)),
                             db=Depends(get_db)):
    grievance_id = validate_uuid(grievance_id, "grievance_id")
    update_data = {
        "manual_resolution": resolution.resolution, "resolution_type": ResolutionType.MANUAL.value,
        "status": GrievanceStatus.RESOLVED.value, "assigned_officer": resolution.officer,
        "updated_at": datetime.now(timezone.utc)}
    if resolution.add_to_service_memory:
        loop2 = asyncio.get_event_loop()
        g = await loop2.run_in_executor(executor, db.grievances.find_one, {"_id": grievance_id})
        if g:
            background_tasks.add_task(add_to_service_memory, g, resolution.resolution, resolution.officer)
    # Officer anomaly detection
    background_tasks.add_task(OfficerAnomalyDetector.check_after_resolution,
                              str(user["_id"]), user["full_name"], resolution.resolution)
    def update():
        return db.grievances.update_one({"_id": grievance_id}, {"$set": update_data})
    loop = asyncio.get_event_loop()
    result = await loop.run_in_executor(executor, update)
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Grievance not found")
    return await get_grievance(grievance_id, user, db)

@app.post("/grievances/{grievance_id}/notes", response_model=GrievanceResponse)
async def add_note(grievance_id: str, note: GrievanceNote,
                   user=Depends(require_role(UserRole.OFFICER.value, UserRole.ADMIN.value)),
                   db=Depends(get_db)):
    grievance_id = validate_uuid(grievance_id, "grievance_id")
    new_note = {"id": str(uuid.uuid4()), "content": note.content, "officer": note.officer,
                "note_type": note.note_type.value, "created_at": datetime.now(timezone.utc)}
    def update():
        return db.grievances.update_one({"_id": grievance_id},
            {"$push": {"notes": new_note}, "$set": {"updated_at": datetime.now(timezone.utc)}})
    loop = asyncio.get_event_loop()
    result = await loop.run_in_executor(executor, update)
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Grievance not found")
    return await get_grievance(grievance_id, user, db)

@app.put("/grievances/{grievance_id}/status", response_model=GrievanceResponse)
async def update_status(grievance_id: str, new_status: GrievanceStatus,
                        user=Depends(require_role(UserRole.OFFICER.value, UserRole.ADMIN.value)),
                        db=Depends(get_db)):
    grievance_id = validate_uuid(grievance_id, "grievance_id")
    loop = asyncio.get_event_loop()
    update_fields: dict = {"status": new_status.value, "updated_at": datetime.now(timezone.utc)}
    if new_status == GrievanceStatus.ESCALATED:
        g = await loop.run_in_executor(executor, db.grievances.find_one, {"_id": grievance_id})
        if g and not g.get("assigned_officer"):
            officer_name = find_officer_for_department(db, g.get("department"))
            if officer_name:
                update_fields["assigned_officer"] = officer_name
    def update():
        return db.grievances.update_one({"_id": grievance_id}, {"$set": update_fields})
    result = await loop.run_in_executor(executor, update)
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Grievance not found")
    return await get_grievance(grievance_id, user, db)

@app.post("/grievances/{grievance_id}/feedback", response_model=GrievanceResponse)
async def add_feedback(grievance_id: str, feedback: ResolutionFeedback,
                       user=Depends(get_current_user), db=Depends(get_db)):
    grievance_id = validate_uuid(grievance_id, "grievance_id")
    # Verify ownership: only the citizen who filed this grievance can leave feedback
    loop = asyncio.get_event_loop()
    g = await loop.run_in_executor(executor, db.grievances.find_one, {"_id": grievance_id})
    if not g:
        raise HTTPException(status_code=404, detail="Grievance not found")
    if user["role"] == UserRole.CITIZEN.value and g.get("citizen_user_id") != str(user["_id"]):
        raise HTTPException(status_code=403, detail="You can only provide feedback on your own grievances")
    def update():
        return db.grievances.update_one({"_id": grievance_id},
            {"$set": {"resolution_feedback": feedback.rating, "updated_at": datetime.now(timezone.utc)}})
    result = await loop.run_in_executor(executor, update)
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Grievance not found")
    return await get_grievance(grievance_id, user, db)

@app.put("/grievances/{grievance_id}/assign", response_model=GrievanceResponse)
async def assign_grievance(grievance_id: str, assignment: GrievanceAssignment,
                           user=Depends(require_role(UserRole.OFFICER.value, UserRole.ADMIN.value)),
                           db=Depends(get_db)):
    grievance_id = validate_uuid(grievance_id, "grievance_id")
    def update():
        return db.grievances.update_one({"_id": grievance_id},
            {"$set": {"assigned_officer": assignment.officer_name,
                      "status": GrievanceStatus.IN_PROGRESS.value,
                      "updated_at": datetime.now(timezone.utc)}})
    loop = asyncio.get_event_loop()
    result = await loop.run_in_executor(executor, update)
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Grievance not found")
    return await get_grievance(grievance_id, user, db)

# -- Citizen confirms self-resolution (auth required) --
class ConfirmResolutionRequest(BaseModel):
    tracking_number: str = Field(..., max_length=30)
    helpful: bool

@app.put("/grievances/confirm-resolution")
@limiter.limit("10/minute")
async def confirm_resolution(request: Request, req: ConfirmResolutionRequest,
                             user=Depends(get_current_user), db=Depends(get_db)):
    tracking_number = sanitize_str(req.tracking_number)
    if not re.match(r'^GRV-\d{4}-[A-Za-z0-9]{6,16}$', tracking_number):
        raise HTTPException(status_code=400, detail="Invalid tracking number format")
    loop = asyncio.get_event_loop()
    g = await loop.run_in_executor(executor, db.grievances.find_one,
                                   {"tracking_number": tracking_number})
    if not g:
        raise HTTPException(status_code=404, detail="Grievance not found")
    # Verify the current user owns this grievance
    if user["role"] == UserRole.CITIZEN.value and g.get("citizen_user_id") != str(user["_id"]):
        raise HTTPException(status_code=403, detail="You can only confirm resolution on your own grievances")
    if g.get("resolution_tier") != ResolutionTier.SELF_RESOLVABLE.value:
        raise HTTPException(status_code=400, detail="This grievance is not eligible for self-resolution")
    if g.get("status") != GrievanceStatus.PENDING.value:
        raise HTTPException(status_code=400, detail="Grievance is no longer pending")

    if req.helpful:
        update_data = {
            "status": GrievanceStatus.RESOLVED.value,
            "resolution_type": ResolutionType.AI.value,
            "updated_at": datetime.now(timezone.utc),
        }
    else:
        # Promote to officer action
        update_data = {
            "resolution_tier": ResolutionTier.OFFICER_ACTION.value,
            "updated_at": datetime.now(timezone.utc),
        }
    await loop.run_in_executor(executor, lambda: db.grievances.update_one(
        {"_id": g["_id"]}, {"$set": update_data}))
    return {"status": "ok", "new_status": update_data.get("status", g["status"]),
            "new_tier": update_data.get("resolution_tier", g.get("resolution_tier"))}

# -- Officer approves / edits AI draft --
class ApproveDraftRequest(BaseModel):
    resolution: Optional[str] = Field(None, max_length=10000)  # None = use AI draft as-is

@app.put("/grievances/{grievance_id}/approve-draft", response_model=GrievanceResponse)
async def approve_draft(grievance_id: str, req: ApproveDraftRequest,
                        background_tasks: BackgroundTasks,
                        user=Depends(require_role(UserRole.OFFICER.value, UserRole.ADMIN.value)),
                        db=Depends(get_db)):
    grievance_id = validate_uuid(grievance_id, "grievance_id")
    loop = asyncio.get_event_loop()
    g = await loop.run_in_executor(executor, db.grievances.find_one, {"_id": grievance_id})
    if not g:
        raise HTTPException(status_code=404, detail="Grievance not found")
    final_text = req.resolution if req.resolution else g.get("ai_resolution", "")
    if not final_text:
        raise HTTPException(status_code=400, detail="No resolution text available")
    update_data = {
        "manual_resolution": final_text,
        "resolution_type": ResolutionType.HYBRID.value,
        "status": GrievanceStatus.RESOLVED.value,
        "assigned_officer": user["full_name"],
        "updated_at": datetime.now(timezone.utc),
    }
    # Also add to service memory
    background_tasks.add_task(add_to_service_memory, g, final_text, user["full_name"])
    # Officer anomaly detection
    background_tasks.add_task(OfficerAnomalyDetector.check_after_resolution,
                              str(user["_id"]), user["full_name"], final_text)
    await loop.run_in_executor(executor, lambda: db.grievances.update_one(
        {"_id": grievance_id}, {"$set": update_data}))
    return await get_grievance(grievance_id, user, db)

# ---------------------------------------------------------------------------
# CHATBOT ENDPOINT (with grievance-filing integration)
# ---------------------------------------------------------------------------
@app.post("/chat", response_model=ChatResponse)
@limiter.limit("15/minute")
async def chat(request: Request, msg: ChatMessage, user=Depends(get_current_user), db=Depends(get_db)):
    lang_name = LANGUAGE_NAMES.get(msg.language.value, "English")
    doc_results = await KnowledgeSearcher.search_documentation(msg.message, limit=3)
    memory_results = await KnowledgeSearcher.search_service_memory(msg.message, limit=3)
    scheme_results = await KnowledgeSearcher.search_schemes(msg.message, limit=3)

    context_parts, sources = [], []
    for d in doc_results:
        context_parts.append(f"Documentation: {d['title']} - {d['content']}")
        sources.append({"type": "documentation", "title": d["title"], "score": d["score"]})
    for m in memory_results:
        context_parts.append(f"Past resolution: {m['query']} -> {m['resolution']}")
        sources.append({"type": "service_memory", "query": m["query"], "score": m["score"]})
    for s in scheme_results:
        context_parts.append(f"Scheme: {s['name']} - {s['description']}. Eligibility: {s['eligibility']}")
        sources.append({"type": "scheme", "name": s["name"], "score": s["score"]})
    context_text = "\n".join(context_parts) if context_parts else "No specific knowledge base context found."

    system_prompt = (
        f"You are Seva, a professional government assistant for the Odisha Panchayati Raj & Drinking Water Department Grievance Portal.\n"
        f"The department handles: rural drinking water (Jal Jeevan Mission, Basudha), MGNREGS, rural housing (PMAY-G), "
        f"panchayat governance, rural livelihoods (NRLM/OLM/SHGs), sanitation (SBM-G), and rural infrastructure (BGBO).\n"
        f"Respond in {lang_name}. Use Markdown formatting.\n\n"
        "RESPONSE RULES (strictly follow):\n"
        "- Keep responses under 120 words. Be concise and direct.\n"
        "- Use **bold** for key terms, names of schemes, and important details.\n"
        "- Use bullet points for lists or steps.\n"
        "- Do not repeat what the citizen just said. Go straight to the answer.\n"
        "- One short greeting max (first message only). After that, be direct.\n"
        "- Do not pad with filler phrases like 'I understand your concern' or 'Thank you for reaching out'.\n\n"
        "GRIEVANCE FILING PROTOCOL:\n"
        "When a citizen wants to file/register/submit a grievance, do NOT file immediately.\n"
        "Instead, gather information through conversation. You need AT LEAST:\n"
        "  1. What is the problem? (specific issue, not just a category)\n"
        "  2. Which DISTRICT? (must be one of the 30 Odisha districts listed below)\n"
        "  3. Where exactly? (village/block/GP name, landmark, or address)\n"
        "  4. When / how long? (when it started or was noticed)\n"
        "  5. Any prior complaints or reference numbers?\n\n"
        "ODISHA DISTRICTS (valid values — you MUST identify which one applies):\n"
        f"  {', '.join(ODISHA_DISTRICTS)}\n\n"
        "After each citizen reply, mentally check which items you still need.\n"
        "- If items are missing, ask for the NEXT missing item (one or two at a time, not all at once).\n"
        "- The district is MANDATORY. If the citizen mentions a village/block/area but not the district, "
        "ask them to confirm which district it falls under.\n"
        "- Once you have items 1-4 at minimum, summarize what you have and ask:\n"
        "  'Shall I file this grievance now, or would you like to add anything?'\n"
        "- ONLY when the citizen confirms (says yes/file/submit/go ahead), include this JSON block "
        "at the END of your response:\n"
        '```json\n{"action":"file_grievance","title":"brief title","description":"full details including location, timing, and specifics","district":"ExactDistrictName"}\n```\n'
        "The description in the JSON must be a comprehensive paragraph combining ALL gathered details.\n"
        "The district in the JSON MUST be one of the 30 valid Odisha district names listed above (exact spelling).\n"
        "NEVER include the JSON block before the citizen confirms filing.\n\n"
        f"Knowledge base context:\n{context_text}"
    )
    messages = [{"role": "system", "content": system_prompt}]
    ALLOWED_ROLES = {"user", "assistant"}
    for h in msg.conversation_history[-10:]:
        role = h.get("role", "user")
        if role not in ALLOWED_ROLES:
            role = "user"
        content = str(h.get("content", ""))[:2000]
        messages.append({"role": role, "content": content})
    messages.append({"role": "user", "content": msg.message})

    try:
        reply = await openai_chat(messages)
        if not reply:
            return ChatResponse(reply="I apologize, I'm having trouble right now. Please try again.", sources=[])

        filed_grievance = None
        json_match = re.search(r'```json\s*(\{.*?"action"\s*:\s*"file_grievance".*?\})\s*```', reply, re.DOTALL)
        if json_match:
            try:
                action_data = json.loads(json_match.group(1))
                if action_data.get("action") == "file_grievance" and action_data.get("title"):
                    district_val = action_data.get("district")
                    if district_val and district_val not in ODISHA_DISTRICTS:
                        district_val = None
                    grievance_data = GrievanceCreate(
                        title=action_data["title"],
                        description=action_data.get("description", action_data["title"]),
                        language=msg.language,
                        district=district_val,
                        citizen_name=user.get("full_name"),
                        citizen_email=user.get("email"),
                    )
                    result = await process_grievance(grievance_data, db, user)
                    filed_grievance = {
                        "tracking_number": result.tracking_number,
                        "status": result.status.value,
                        "department": result.department.value,
                        "priority": result.priority.value,
                    }
                    reply = reply[:json_match.start()].strip()
            except Exception as e:
                logger.error("Chat grievance filing error: %s", e)

        return ChatResponse(reply=reply, sources=sources, filed_grievance=filed_grievance)
    except Exception as e:
        logger.error("Chat error: %s", e)
        return ChatResponse(reply="I apologize, I'm having trouble right now. Please try again.", sources=[])

# ---------------------------------------------------------------------------
# KNOWLEDGE BASE ENDPOINTS
# ---------------------------------------------------------------------------
async def add_to_service_memory(grievance: dict, resolution: str, officer: str):
    try:
        query = f"{grievance['title']} {grievance['description']}"
        embedding = await get_openai_embedding(query)
        point = PointStruct(id=str(uuid.uuid4()), vector=embedding,
            payload={"query": query, "resolution": resolution,
                     "category": grievance.get("department", "general"),
                     "agent_name": officer,
                     "tracking_number": grievance.get("tracking_number"),
                     "grievance_id": str(grievance.get("_id", "")),
                     "created_at": datetime.now(timezone.utc).isoformat()})
        qdrant.upsert(collection_name="service_memory", points=[point], wait=True)
    except Exception as e:
        logger.error("Error adding to service memory: %s", e)

@app.post("/knowledge/documentation")
async def add_documentation(entry: KnowledgeEntry,
                            user=Depends(require_role(UserRole.OFFICER.value, UserRole.ADMIN.value))):
    if not qdrant:
        raise HTTPException(status_code=503, detail="Knowledge base (Qdrant) is not available")
    embedding = await get_openai_embedding(entry.content)
    point = PointStruct(id=str(uuid.uuid4()), vector=embedding,
        payload={"title": entry.title, "content": entry.content,
                 "category": entry.category.value, "created_at": datetime.now(timezone.utc).isoformat()})
    qdrant.upsert(collection_name="documentation", points=[point], wait=True)
    return {"message": "Documentation added successfully"}

@app.post("/knowledge/service-memory")
async def add_service_memory_endpoint(entry: ServiceMemoryEntry,
                                       user=Depends(require_role(UserRole.OFFICER.value, UserRole.ADMIN.value))):
    if not qdrant:
        raise HTTPException(status_code=503, detail="Knowledge base (Qdrant) is not available")
    embedding = await get_openai_embedding(entry.query)
    point = PointStruct(id=str(uuid.uuid4()), vector=embedding,
        payload={"query": entry.query, "resolution": entry.resolution,
                 "category": entry.category.value, "agent_name": entry.agent_name,
                 "created_at": datetime.now(timezone.utc).isoformat()})
    qdrant.upsert(collection_name="service_memory", points=[point], wait=True)
    return {"message": "Service memory added successfully"}

@app.post("/knowledge/schemes")
async def add_scheme(entry: SchemeEntry,
                     user=Depends(require_role(UserRole.OFFICER.value, UserRole.ADMIN.value))):
    if not qdrant:
        raise HTTPException(status_code=503, detail="Knowledge base (Qdrant) is not available")
    text = f"{entry.name} {entry.description} {entry.eligibility} {entry.how_to_apply}"
    embedding = await get_openai_embedding(text)
    point = PointStruct(id=str(uuid.uuid4()), vector=embedding,
        payload={"name": entry.name, "description": entry.description,
                 "eligibility": entry.eligibility, "department": entry.department.value,
                 "how_to_apply": entry.how_to_apply,
                 "eligibility_questions": entry.eligibility_questions,
                 "documents": [],
                 "created_at": datetime.now(timezone.utc).isoformat()})
    qdrant.upsert(collection_name="schemes", points=[point], wait=True)
    return {"message": "Scheme added successfully"}

@app.post("/knowledge/schemes/generate-questions")
async def generate_questions_preview(req: GenerateQuestionsRequest,
                                      user=Depends(require_role(UserRole.OFFICER.value, UserRole.ADMIN.value))):
    questions = await generate_eligibility_questions(req.eligibility, req.scheme_name)
    return {"questions": questions}

@app.post("/knowledge/schemes/backfill-questions")
async def backfill_scheme_questions(
    user=Depends(require_role(UserRole.OFFICER.value, UserRole.ADMIN.value))):
    if not qdrant:
        raise HTTPException(status_code=503, detail="Knowledge base (Qdrant) is not available")
    zero_vector = [0.0] * EMBEDDING_DIM
    results = qdrant.query_points(collection_name="schemes", query=zero_vector,
                                   limit=100, with_payload=True, with_vectors=False)
    updated = 0
    for p in results.points:
        existing = p.payload.get("eligibility_questions", [])
        if not existing:
            elig = p.payload.get("eligibility", "")
            name = p.payload.get("name", "")
            if elig:
                questions = await generate_eligibility_questions(elig, name)
                if questions:
                    qdrant.set_payload(collection_name="schemes",
                                       payload={"eligibility_questions": questions},
                                       points=[p.id])
                    updated += 1
    return {"message": f"Backfilled questions for {updated} schemes"}

@app.get("/knowledge/schemes/{scheme_id}", response_model=SchemeResponse)
async def get_scheme_detail(scheme_id: str):
    if not qdrant:
        raise HTTPException(status_code=503, detail="Knowledge base (Qdrant) is not available")
    points = qdrant.retrieve(collection_name="schemes", ids=[scheme_id], with_payload=True)
    if not points:
        raise HTTPException(status_code=404, detail="Scheme not found")
    p = points[0]
    return {"id": p.id, "name": p.payload.get("name", ""), "description": p.payload.get("description", ""),
            "eligibility": p.payload.get("eligibility", ""), "department": p.payload.get("department", ""),
            "how_to_apply": p.payload.get("how_to_apply", ""), "created_at": p.payload.get("created_at", ""),
            "eligibility_questions": p.payload.get("eligibility_questions", []),
            "documents": p.payload.get("documents", [])}

@app.put("/knowledge/schemes/{scheme_id}/questions")
async def update_scheme_questions(scheme_id: str, req: UpdateQuestionsRequest,
                                   user=Depends(require_role(UserRole.OFFICER.value, UserRole.ADMIN.value))):
    if not qdrant:
        raise HTTPException(status_code=503, detail="Knowledge base (Qdrant) is not available")
    points = qdrant.retrieve(collection_name="schemes", ids=[scheme_id], with_payload=True)
    if not points:
        raise HTTPException(status_code=404, detail="Scheme not found")
    qdrant.set_payload(collection_name="schemes",
                       payload={"eligibility_questions": req.questions},
                       points=[scheme_id])
    return {"message": "Questions updated successfully"}

@app.get("/knowledge/documentation", response_model=List[DocumentationResponse])
async def get_documentation(category: Optional[Department] = None,
                            limit: int = Query(50, ge=1, le=100), offset: int = Query(0, ge=0, le=10000),
                            user=Depends(require_role(UserRole.OFFICER.value, UserRole.ADMIN.value))):
    if not qdrant:
        return []
    filt = Filter(must=[FieldCondition(key="category", match=MatchValue(value=category.value))]) if category else None
    zero_vector = [0.0] * EMBEDDING_DIM
    results = qdrant.query_points(collection_name="documentation", query=zero_vector,
                            query_filter=filt, limit=limit, offset=offset, with_payload=True, with_vectors=False)
    return [{"id": p.id, "title": p.payload.get("title", ""), "content": p.payload.get("content", ""),
             "category": p.payload.get("category", ""), "created_at": p.payload.get("created_at", "")} for p in results.points]

@app.get("/knowledge/service-memory", response_model=List[ServiceMemoryResponse])
async def get_service_memory(category: Optional[Department] = None,
                             limit: int = Query(50, ge=1, le=100), offset: int = Query(0, ge=0, le=10000),
                             user=Depends(require_role(UserRole.OFFICER.value, UserRole.ADMIN.value))):
    if not qdrant:
        return []
    filt = Filter(must=[FieldCondition(key="category", match=MatchValue(value=category.value))]) if category else None
    zero_vector = [0.0] * EMBEDDING_DIM
    results = qdrant.query_points(collection_name="service_memory", query=zero_vector,
                            query_filter=filt, limit=limit, offset=offset, with_payload=True, with_vectors=False)
    return [{"id": p.id, "query": p.payload.get("query", ""), "resolution": p.payload.get("resolution", ""),
             "category": p.payload.get("category", ""), "agent_name": p.payload.get("agent_name", ""),
             "tracking_number": p.payload.get("tracking_number"),
             "grievance_id": p.payload.get("grievance_id"),
             "created_at": p.payload.get("created_at", "")} for p in results.points]

@app.delete("/knowledge/documentation/{doc_id}")
async def delete_documentation(doc_id: str,
                                user=Depends(require_role(UserRole.OFFICER.value, UserRole.ADMIN.value))):
    if not qdrant:
        raise HTTPException(status_code=503, detail="Knowledge base (Qdrant) is not available")
    qdrant.delete(collection_name="documentation", points_selector=PointIdsList(points=[doc_id]))
    return {"message": "Documentation deleted"}

@app.delete("/knowledge/service-memory/{mem_id}")
async def delete_service_memory(mem_id: str,
                                 user=Depends(require_role(UserRole.OFFICER.value, UserRole.ADMIN.value))):
    if not qdrant:
        raise HTTPException(status_code=503, detail="Knowledge base (Qdrant) is not available")
    qdrant.delete(collection_name="service_memory", points_selector=PointIdsList(points=[mem_id]))
    return {"message": "Service memory entry deleted"}

@app.delete("/knowledge/schemes/{scheme_id}")
async def delete_scheme(scheme_id: str,
                         user=Depends(require_role(UserRole.OFFICER.value, UserRole.ADMIN.value))):
    if not qdrant:
        raise HTTPException(status_code=503, detail="Knowledge base (Qdrant) is not available")
    qdrant.delete(collection_name="schemes", points_selector=PointIdsList(points=[scheme_id]))
    return {"message": "Scheme deleted"}

@app.put("/knowledge/schemes/{scheme_id}")
async def update_scheme(scheme_id: str, entry: SchemeEntry,
                         user=Depends(require_role(UserRole.OFFICER.value, UserRole.ADMIN.value))):
    if not qdrant:
        raise HTTPException(status_code=503, detail="Knowledge base (Qdrant) is not available")
    points = qdrant.retrieve(collection_name="schemes", ids=[scheme_id], with_payload=True)
    if not points:
        raise HTTPException(status_code=404, detail="Scheme not found")
    # Re-embed with updated text
    text = f"{entry.name} {entry.description} {entry.eligibility} {entry.how_to_apply}"
    embedding = await get_openai_embedding(text)
    # Preserve existing questions if not provided in the update
    existing_questions = points[0].payload.get("eligibility_questions", [])
    new_questions = entry.eligibility_questions if entry.eligibility_questions else existing_questions
    # Preserve existing documents
    existing_documents = points[0].payload.get("documents", [])
    point = PointStruct(id=scheme_id, vector=embedding,
        payload={"name": entry.name, "description": entry.description,
                 "eligibility": entry.eligibility, "department": entry.department.value,
                 "how_to_apply": entry.how_to_apply,
                 "eligibility_questions": new_questions,
                 "documents": existing_documents,
                 "created_at": points[0].payload.get("created_at", datetime.now(timezone.utc).isoformat())})
    qdrant.upsert(collection_name="schemes", points=[point], wait=True)
    return {"message": "Scheme updated successfully"}

@app.get("/knowledge/schemes", response_model=List[SchemeResponse])
async def get_schemes(department: Optional[Department] = None,
                      limit: int = Query(50, ge=1, le=100), offset: int = Query(0, ge=0, le=10000)):
    if not qdrant:
        return []
    filt = Filter(must=[FieldCondition(key="department", match=MatchValue(value=department.value))]) if department else None
    zero_vector = [0.0] * EMBEDDING_DIM
    results = qdrant.query_points(collection_name="schemes", query=zero_vector,
                            query_filter=filt, limit=limit, offset=offset, with_payload=True, with_vectors=False)
    schemes = [{"id": p.id, "name": p.payload.get("name", ""), "description": p.payload.get("description", ""),
             "eligibility": p.payload.get("eligibility", ""), "department": p.payload.get("department", ""),
             "how_to_apply": p.payload.get("how_to_apply", ""), "created_at": p.payload.get("created_at", ""),
             "eligibility_questions": p.payload.get("eligibility_questions", []),
             "documents": p.payload.get("documents", [])} for p in results.points]
    schemes.sort(key=lambda s: s.get("created_at", ""), reverse=True)
    return schemes

@app.post("/knowledge/schemes/search")
@limiter.limit("20/minute")
async def search_schemes_endpoint(request: Request, query: str = Query(...)):
    return await KnowledgeSearcher.search_schemes(query, limit=10)

# ---------------------------------------------------------------------------
# SCHEME DOCUMENT PROCESSING
# ---------------------------------------------------------------------------

def _extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF using PyPDF2. Returns empty string for scanned/image PDFs."""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text.strip())
        return "\n\n".join(t for t in text_parts if t)
    except Exception as e:
        logger.error("PDF text extraction error: %s", e)
        return ""


def _extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file including paragraphs and tables."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)
    except Exception as e:
        logger.error("DOCX text extraction error: %s", e)
        return ""


def _image_to_base64(file_bytes: bytes, content_type: str) -> str:
    """Encode image bytes to base64 data URL for the OpenAI Vision API."""
    b64 = base64.b64encode(file_bytes).decode("utf-8")
    media = content_type or "image/png"
    return f"data:{media};base64,{b64}"


def _pdf_pages_to_images(file_bytes: bytes) -> list:
    """Convert first few PDF pages to base64 images for Vision API (scanned PDFs)."""
    images = []
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        # For scanned PDFs, send the raw bytes as a single image if possible
        # PyPDF2 can't render pages to images, so we send the PDF bytes directly
        # and let the Vision API handle it. For truly image-based PDFs, we
        # extract embedded images from the first few pages.
        for page_num, page in enumerate(reader.pages[:5]):
            if "/XObject" in (page.get("/Resources") or {}):
                x_objects = page["/Resources"]["/XObject"].get_object()
                for obj_name in x_objects:
                    obj = x_objects[obj_name].get_object()
                    if obj.get("/Subtype") == "/Image":
                        try:
                            data = obj.get_data()
                            if len(data) > 100:  # skip tiny images (icons, etc)
                                b64 = base64.b64encode(data).decode("utf-8")
                                # Detect format from filter
                                filt = obj.get("/Filter", "")
                                if isinstance(filt, list):
                                    filt = filt[0] if filt else ""
                                filt_str = str(filt)
                                if "DCT" in filt_str:
                                    media = "image/jpeg"
                                elif "PNG" in filt_str or "Flat" in filt_str:
                                    media = "image/png"
                                else:
                                    media = "image/jpeg"
                                images.append(f"data:{media};base64,{b64}")
                                if len(images) >= 3:
                                    break
                        except Exception:
                            continue
            if len(images) >= 3:
                break
    except Exception as e:
        logger.error("PDF image extraction error: %s", e)
    return images


async def _extract_scheme_with_vision(image_data_urls: list) -> Optional[dict]:
    """Use OpenAI Vision API to extract scheme info from images."""
    content = [
        {"type": "text", "text": (
            "Analyze these images. Determine if they describe a government welfare scheme, "
            "public benefit program, or official government policy. "
            "Return JSON: {\"relevant\": true/false, \"relevance_reason\": \"why it is or isn't a scheme doc\", "
            "\"name\": \"...\", \"description\": \"...\", "
            "\"eligibility\": \"...\", \"department\": \"...\", \"how_to_apply\": \"...\"}. "
            "If not relevant, set all fields to empty strings."
        )}
    ]
    for url in image_data_urls[:3]:
        content.append({"type": "image_url", "image_url": {"url": url, "detail": "high"}})
    try:
        resp = await openai_client.chat.completions.create(
            model=OPENAI_VISION_MODEL,
            messages=[{"role": "user", "content": content}],
            response_format={"type": "json_object"},
            max_tokens=2000
        )
        result = resp.choices[0].message.content.strip()
        return json.loads(result)
    except Exception as e:
        logger.error("Vision API extraction error: %s", e)
        return None


async def _extract_scheme_from_text(text: str) -> Optional[dict]:
    """Use OpenAI to extract structured scheme info from text."""
    prompt = (
        "Analyze the following text. First determine if this document describes a "
        "government welfare scheme, public benefit program, subsidy, grant, or official "
        "government policy/service meant for citizens or organizations.\n\n"
        "Return JSON only:\n"
        '{"relevant": true/false, '
        '"relevance_reason": "one-line explanation of why it is or is not a scheme document", '
        '"name": "scheme name", "description": "comprehensive description", '
        '"eligibility": "detailed eligibility criteria", '
        '"department": "responsible department", '
        '"how_to_apply": "application process"}\n\n'
        "If relevant is false, set name/description/eligibility/department/how_to_apply to empty strings.\n"
        "If relevant is true but a field is not found, use an empty string for that field.\n"
        "For department, try to match one of: panchayati_raj, rural_development, "
        "rural_water_supply, sanitation, rural_housing, rural_roads, "
        "social_welfare, tribal_welfare, agriculture, education, health, "
        "womens_welfare, disability_welfare, revenue, general.\n\n"
        f"Document text:\n{truncate_text(text, 6000)}"
    )
    try:
        result = await openai_chat([{"role": "user", "content": prompt}], json_mode=True)
        if not result:
            return None
        return json.loads(result)
    except Exception as e:
        logger.error("Text extraction error: %s", e)
        return None


@app.post("/knowledge/schemes/process-documents")
@limiter.limit("10/minute")
async def process_scheme_documents(
    request: Request,
    files: List[UploadFile] = File(...),
    user=Depends(require_role(UserRole.OFFICER.value, UserRole.ADMIN.value))
):
    """Process uploaded scheme documents: extract text, send to OpenAI, store in GridFS."""
    if len(files) > SCHEME_MAX_FILES:
        raise HTTPException(400, f"Maximum {SCHEME_MAX_FILES} files allowed")

    all_text = []
    vision_images = []
    stored_docs = []
    loop = asyncio.get_event_loop()

    for f in files:
        content_type = f.content_type or ""
        if content_type not in SCHEME_ALLOWED_MIMES:
            raise HTTPException(400, f"Unsupported file type: {f.filename}. Use PDF, DOCX, PNG, or JPG.")
        file_bytes = await f.read()
        if len(file_bytes) > SCHEME_MAX_FILE_SIZE:
            raise HTTPException(400, f"{f.filename} exceeds 10MB limit")

        # Store in GridFS
        try:
            file_id = await loop.run_in_executor(
                executor,
                lambda fb=file_bytes, fn=f.filename, ct=content_type: gfs.put(
                    fb, filename=fn, content_type=ct, metadata={"type": "scheme_document"}
                )
            )
            stored_docs.append({
                "id": str(file_id),
                "filename": f.filename,
                "content_type": content_type,
                "size": len(file_bytes)
            })
        except Exception as e:
            logger.error("GridFS storage error for %s: %s", f.filename, e)
            raise HTTPException(500, f"Failed to store {f.filename}")

        # Extract text based on file type
        if content_type == "application/pdf":
            text = await loop.run_in_executor(executor, _extract_text_from_pdf, file_bytes)
            if text and len(text.strip()) > 50:
                all_text.append(text)
            else:
                # Scanned PDF — try extracting embedded images for Vision API
                images = await loop.run_in_executor(executor, _pdf_pages_to_images, file_bytes)
                if images:
                    vision_images.extend(images)
                else:
                    # Last resort: encode entire PDF first page as generic image
                    logger.warning("Scanned PDF %s: no extractable text or images", f.filename)
        elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            text = await loop.run_in_executor(executor, _extract_text_from_docx, file_bytes)
            if text:
                all_text.append(text)
        elif content_type.startswith("image/"):
            data_url = _image_to_base64(file_bytes, content_type)
            vision_images.append(data_url)

    # Extract scheme info
    extracted = None
    combined_text = "\n\n---\n\n".join(all_text)

    if combined_text.strip():
        extracted = await _extract_scheme_from_text(combined_text)

    if not extracted and vision_images:
        extracted = await _extract_scheme_with_vision(vision_images)

    if not extracted:
        # Return stored docs even if extraction failed
        return {
            "extracted": {"name": "", "description": "", "eligibility": "",
                         "department": "", "how_to_apply": ""},
            "documents": stored_docs,
            "eligibility_questions": [],
            "relevant": False,
            "relevance_reason": "Could not extract any readable text or information from the uploaded files.",
            "message": "Could not extract any details from these files."
        }

    # Check relevance flag from AI
    is_relevant = extracted.get("relevant", True)  # default True for backward compat
    relevance_reason = extracted.get("relevance_reason", "")

    # Check if all key fields are empty (another signal of irrelevance)
    key_fields = ["name", "description", "eligibility", "how_to_apply"]
    all_empty = all(not extracted.get(f, "").strip() for f in key_fields)

    if not is_relevant or all_empty:
        reason = relevance_reason or "The uploaded files don't contain information about a government scheme."
        return {
            "extracted": {"name": "", "description": "", "eligibility": "",
                         "department": "", "how_to_apply": ""},
            "documents": stored_docs,
            "eligibility_questions": [],
            "relevant": False,
            "relevance_reason": reason,
            "message": reason
        }

    # Clean out the relevance metadata before sending extracted data to frontend
    extracted.pop("relevant", None)
    extracted.pop("relevance_reason", None)

    # Generate eligibility questions from extracted eligibility text
    eligibility_questions = []
    elig_text = extracted.get("eligibility", "")
    scheme_name = extracted.get("name", "")
    if elig_text:
        eligibility_questions = await generate_eligibility_questions(elig_text, scheme_name)

    return {
        "extracted": extracted,
        "documents": stored_docs,
        "eligibility_questions": eligibility_questions,
        "relevant": True,
        "message": "Extracted — review and edit below"
    }


@app.put("/knowledge/schemes/{scheme_id}/documents")
async def update_scheme_documents(scheme_id: str,
                                   request: Request,
                                   user=Depends(require_role(UserRole.OFFICER.value, UserRole.ADMIN.value))):
    """Update the documents list for an existing scheme."""
    if not qdrant:
        raise HTTPException(status_code=503, detail="Knowledge base (Qdrant) is not available")
    body = await request.json()
    documents = body.get("documents", [])
    points = qdrant.retrieve(collection_name="schemes", ids=[scheme_id], with_payload=True)
    if not points:
        raise HTTPException(status_code=404, detail="Scheme not found")
    qdrant.set_payload(collection_name="schemes",
                       payload={"documents": documents},
                       points=[scheme_id])
    return {"message": "Documents updated"}


# ---------------------------------------------------------------------------
# ANALYTICS ENDPOINTS (real computation)
# ---------------------------------------------------------------------------
@app.get("/analytics", response_model=AnalyticsResponse)
async def get_analytics(days: int = Query(30, ge=1, le=365),
                        department: Optional[str] = Query(None),
                        user=Depends(require_role(UserRole.OFFICER.value, UserRole.ADMIN.value)),
                        db=Depends(get_db)):
    start_date = datetime.now(timezone.utc) - timedelta(days=days)
    base: dict = {"created_at": {"$gte": start_date}}
    if department:
        base["department"] = department
    def fetch():
        total = db.grievances.count_documents(base)
        self_resolved = db.grievances.count_documents({
            **base, "status": "resolved",
            "resolution_tier": "self_resolvable", "resolution_type": "ai"})
        ai_drafted = db.grievances.count_documents({
            **base, "ai_resolution": {"$exists": True, "$ne": None}})
        escalated = db.grievances.count_documents({**base, "status": "escalated"})
        sent_results = list(db.grievances.aggregate([
            {"$match": base},
            {"$group": {"_id": "$sentiment", "count": {"$sum": 1}}}]))
        dept_results = list(db.grievances.aggregate([
            {"$match": base},
            {"$group": {"_id": "$department", "count": {"$sum": 1}}}]))
        # Real avg resolution time
        avg_pipe = [
            {"$match": {**base, "status": "resolved"}},
            {"$project": {"duration": {"$subtract": ["$updated_at", "$created_at"]}}},
            {"$group": {"_id": None, "avg_ms": {"$avg": "$duration"}}}]
        avg_result = list(db.grievances.aggregate(avg_pipe))
        avg_hours = (avg_result[0]["avg_ms"] / 3600000) if avg_result and avg_result[0].get("avg_ms") else 0.0
        # Pending and SLA breached counts
        now = datetime.now(timezone.utc)
        pending_count = db.grievances.count_documents({
            **base, "status": {"$in": ["pending", "in_progress", "escalated"]}})
        sla_breached_count = db.grievances.count_documents({
            **base, "status": {"$in": ["pending", "in_progress", "escalated"]},
            "sla_deadline": {"$lt": now}})
        # Deadline alerts (near-breach: within 25% of remaining time)
        deadline_alert_count = db.grievances.count_documents({
            **base, "status": {"$in": ["pending", "in_progress", "escalated"]},
            "$or": [
                {"sla_deadline": {"$lt": now}},
                {"estimated_resolution_deadline": {"$lt": now, "$ne": None}},
            ]})
        # Status distribution
        status_results = list(db.grievances.aggregate([
            {"$match": base},
            {"$group": {"_id": "$status", "count": {"$sum": 1}}}]))
        status_dist = {s["_id"]: s["count"] for s in status_results}
        # Top districts by grievance count (all, not capped at 10)
        district_pipe = [
            {"$match": {**base, "district": {"$ne": None}}},
            {"$group": {"_id": "$district", "count": {"$sum": 1}}},
            {"$sort": {"count": -1}}]
        district_results = list(db.grievances.aggregate(district_pipe))
        top_districts = [{"district": d["_id"], "count": d["count"]} for d in district_results]
        # Top complaint titles
        complaint_pipe = [
            {"$match": base},
            {"$group": {"_id": "$title", "department": {"$first": "$department"}, "count": {"$sum": 1}}},
            {"$sort": {"count": -1}}, {"$limit": 10}]
        complaint_results = list(db.grievances.aggregate(complaint_pipe))
        top_complaints = [{"title": c["_id"], "department": c.get("department", "general"),
                           "count": c["count"]} for c in complaint_results]
        return {"total": total, "self_resolved": self_resolved, "ai_drafted": ai_drafted,
                "escalated": escalated,
                "sent": sent_results, "dept": dept_results, "avg_hours": avg_hours,
                "pending_count": pending_count, "sla_breached_count": sla_breached_count,
                "deadline_alert_count": deadline_alert_count,
                "status_dist": status_dist,
                "top_districts": top_districts, "top_complaints": top_complaints}
    loop = asyncio.get_event_loop()
    data = await loop.run_in_executor(executor, fetch)
    return AnalyticsResponse(
        total_grievances=data["total"], self_resolved=data["self_resolved"],
        ai_drafted=data["ai_drafted"],
        escalated_to_human=data["escalated"], avg_resolution_time=round(data["avg_hours"], 1),
        pending_count=data["pending_count"], sla_breached_count=data["sla_breached_count"],
        deadline_alerts=data.get("deadline_alert_count", 0),
        status_distribution=data["status_dist"],
        sentiment_distribution={i["_id"]: i["count"] for i in data["sent"]},
        department_distribution={i["_id"]: i["count"] for i in data["dept"]},
        top_districts=data["top_districts"],
        top_complaints=data["top_complaints"])

@app.get("/analytics/geographical", response_model=GeoAnalyticsResponse)
async def get_geo_analytics(days: int = Query(30, ge=1, le=365),
                            user=Depends(require_role(UserRole.OFFICER.value, UserRole.ADMIN.value)),
                            db=Depends(get_db)):
    start_date = datetime.now(timezone.utc) - timedelta(days=days)
    def fetch():
        hotspots, district_counts = [], {}
        for g in db.grievances.find({"created_at": {"$gte": start_date}, "location": {"$exists": True, "$ne": None}}).limit(200):
            coords = g.get("location", {}).get("coordinates", [])
            if len(coords) == 2:
                hotspots.append({"latitude": coords[1], "longitude": coords[0],
                                 "priority": g.get("priority"), "department": g.get("department")})
            d = g.get("district", "unknown")
            district_counts[d] = district_counts.get(d, 0) + 1
        return {"hotspots": hotspots, "district_counts": district_counts}
    loop = asyncio.get_event_loop()
    data = await loop.run_in_executor(executor, fetch)
    return GeoAnalyticsResponse(hotspot_coordinates=data["hotspots"],
                                region_distribution={}, district_distribution=data["district_counts"])

# ---------------------------------------------------------------------------
# ATTACHMENT ENDPOINTS
# ---------------------------------------------------------------------------
@app.get("/grievances/{grievance_id}/attachments")
async def get_attachments(grievance_id: str, user=Depends(get_current_user), db=Depends(get_db)):
    grievance_id = validate_uuid(grievance_id, "grievance_id")
    loop = asyncio.get_event_loop()
    g = await loop.run_in_executor(executor, db.grievances.find_one, {"_id": grievance_id})
    if not g:
        raise HTTPException(status_code=404, detail="Grievance not found")
    att_ids = g.get("attachments", [])
    result = []
    for aid in att_ids:
        try:
            fobj = await loop.run_in_executor(executor, lambda a=aid: gfs.get(ObjectId(a)))
            result.append({"id": aid, "filename": fobj.filename, "content_type": fobj.content_type,
                           "length": fobj.length})
        except Exception:
            pass
    return result

@app.get("/attachments/{attachment_id}")
async def get_attachment(attachment_id: str):
    try:
        loop = asyncio.get_event_loop()
        fobj = await loop.run_in_executor(executor, lambda: gfs.get(ObjectId(attachment_id)))
        def iterfile():
            while True:
                chunk = fobj.read(8192)
                if not chunk:
                    break
                yield chunk
        return StreamingResponse(iterfile(), media_type=fobj.content_type or "application/octet-stream",
                                 headers={"Content-Disposition": f'inline; filename="{fobj.filename}"'})
    except Exception:
        raise HTTPException(status_code=404, detail="Attachment not found")

# ---------------------------------------------------------------------------
# IP GEOLOCATION ENDPOINT
# ---------------------------------------------------------------------------
@app.get("/location/from-ip")
@limiter.limit("5/minute")
async def location_from_ip(request: Request):
    ip = request.headers.get("X-Forwarded-For", "").split(",")[0].strip()
    if not ip:
        ip = request.headers.get("X-Real-IP", "")
    if not ip and request.client:
        ip = request.client.host
    if not ip:
        raise HTTPException(status_code=400, detail="Could not determine IP address")
    try:
        addr = ipaddress.ip_address(ip)
        if addr.is_private or addr.is_loopback:
            raise HTTPException(status_code=400, detail="Private/localhost IP — use browser geolocation instead")
    except ValueError:
        pass
    try:
        async with httpx.AsyncClient(timeout=5.0) as client:
            resp = await client.get(f"http://ip-api.com/json/{ip}",
                                    params={"fields": "status,lat,lon,city,regionName,country"})
            data = resp.json()
            if data.get("status") != "success":
                raise HTTPException(status_code=400, detail="IP geolocation failed")
            return {"latitude": data.get("lat"), "longitude": data.get("lon"),
                    "city": data.get("city"), "region": data.get("regionName"),
                    "country": data.get("country")}
    except HTTPException:
        raise
    except Exception as e:
        logger.error("IP geolocation error: %s", e)
        raise HTTPException(status_code=500, detail="Geolocation service error")

# ---------------------------------------------------------------------------
# SPAM / PHOTO ID ENDPOINTS
# ---------------------------------------------------------------------------
@app.post("/auth/upload-photo-id")
async def upload_photo_id(request: Request, file: UploadFile = File(...),
                          user=Depends(get_current_user), db=Depends(get_db)):
    if user["role"] != UserRole.CITIZEN.value:
        raise HTTPException(status_code=403, detail="Only citizens can upload photo ID")
    if file.content_type not in ("image/jpeg", "image/png", "image/webp"):
        raise HTTPException(status_code=400, detail="Only JPEG, PNG, or WebP images allowed")
    content = await file.read()
    if len(content) > 10 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File exceeds 10MB limit")
    loop = asyncio.get_event_loop()
    fid = await loop.run_in_executor(executor, lambda: gfs.put(
        content, filename=f"photo_id_{user['_id']}", content_type=file.content_type))
    await loop.run_in_executor(executor, lambda: db.spam_tracking.update_one(
        {"_id": str(user["_id"])},
        {"$set": {"photo_id_file_id": str(fid), "photo_id_status": "pending_review"}},
        upsert=True))
    return {"detail": "Photo ID uploaded for review", "status": "pending_review"}

@app.get("/admin/spam-flagged")
async def get_spam_flagged(user=Depends(require_role(UserRole.ADMIN.value)), db=Depends(get_db)):
    loop = asyncio.get_event_loop()
    records = await loop.run_in_executor(executor, lambda: list(
        db.spam_tracking.find({"is_blocked": True}).limit(50)))
    result = []
    for r in records:
        u = await loop.run_in_executor(executor, db.users.find_one, {"_id": r["_id"]})
        result.append({
            "user_id": r["_id"],
            "username": u["username"] if u else "unknown",
            "full_name": u["full_name"] if u else "unknown",
            "spam_score": r.get("spam_score", 0),
            "blocked_at": r.get("blocked_at"),
            "photo_id_status": r.get("photo_id_status", "none"),
            "photo_id_file_id": r.get("photo_id_file_id"),
            "pattern_flags": r.get("pattern_flags", [])[-5:],
        })
    return result

@app.get("/admin/spam-flagged/{user_id}")
async def get_spam_flag_detail(user_id: str,
                               user=Depends(require_role(UserRole.OFFICER.value, UserRole.ADMIN.value)),
                               db=Depends(get_db)):
    user_id = validate_uuid(user_id, "user_id")
    loop = asyncio.get_event_loop()
    record = await loop.run_in_executor(executor, db.spam_tracking.find_one, {"_id": user_id})
    if not record:
        raise HTTPException(status_code=404, detail="No spam record found for this user")
    u = await loop.run_in_executor(executor, db.users.find_one, {"_id": user_id})
    grievances = await loop.run_in_executor(executor, lambda: list(
        db.grievances.find({"citizen_user_id": user_id}).sort("created_at", -1).limit(20)))
    # Build grievance list with content hashes
    grv_list = []
    texts = []
    for g in grievances:
        text = f"{g.get('title', '')} {g.get('description', '')}".lower().strip()
        texts.append(text)
        grv_list.append({
            "id": g["_id"],
            "tracking_number": g.get("tracking_number"),
            "title": g.get("title"),
            "description": g.get("description", "")[:200],
            "status": g.get("status"),
            "department": g.get("department"),
            "priority": g.get("priority"),
            "created_at": g.get("created_at"),
            "content_hash": hashlib.sha256(
                f"{g.get('title', '')}|{g.get('description', '')}".lower().strip().encode()
            ).hexdigest(),
            "similar_to": [],
        })
    # Pairwise near-duplicate detection (SequenceMatcher, threshold 75%)
    for i in range(len(grv_list)):
        for j in range(i + 1, len(grv_list)):
            ratio = difflib.SequenceMatcher(None, texts[i], texts[j]).ratio()
            if ratio >= 0.75:
                pct = round(ratio * 100)
                grv_list[i]["similar_to"].append(
                    {"tracking_number": grv_list[j]["tracking_number"], "similarity": pct})
                grv_list[j]["similar_to"].append(
                    {"tracking_number": grv_list[i]["tracking_number"], "similarity": pct})
    return {
        "user_id": user_id,
        "username": u["username"] if u else "unknown",
        "full_name": u["full_name"] if u else "unknown",
        "email": u.get("email") if u else None,
        "phone": u.get("phone") if u else None,
        "spam_score": record.get("spam_score", 0),
        "is_blocked": record.get("is_blocked", False),
        "blocked_at": record.get("blocked_at"),
        "photo_id_status": record.get("photo_id_status", "none"),
        "photo_id_file_id": record.get("photo_id_file_id"),
        "pattern_flags": record.get("pattern_flags", []),
        "ip_addresses": record.get("ip_addresses", []),
        "filing_timestamps": record.get("filing_timestamps", []),
        "duplicate_hashes": record.get("duplicate_hashes", []),
        "grievances": grv_list,
    }

@app.put("/admin/spam-flagged/{user_id}/review")
async def review_spam_user(user_id: str, approved: bool = Query(...),
                           user=Depends(require_role(UserRole.ADMIN.value)), db=Depends(get_db)):
    user_id = validate_uuid(user_id, "user_id")
    loop = asyncio.get_event_loop()
    if approved:
        await loop.run_in_executor(executor, lambda: db.spam_tracking.update_one(
            {"_id": user_id},
            {"$set": {"is_blocked": False, "spam_score": 0.0, "photo_id_status": "approved",
                      "pattern_flags": []}}))
        return {"detail": "User unblocked and verified"}
    else:
        await loop.run_in_executor(executor, lambda: db.spam_tracking.update_one(
            {"_id": user_id}, {"$set": {"photo_id_status": "rejected"}}))
        return {"detail": "Photo ID rejected, user remains blocked"}

@app.get("/auth/spam-status")
async def get_spam_status(user=Depends(get_current_user), db=Depends(get_db)):
    loop = asyncio.get_event_loop()
    record = await loop.run_in_executor(executor, db.spam_tracking.find_one, {"_id": str(user["_id"])})
    if not record:
        return {"is_blocked": False, "photo_id_status": "none"}
    return {"is_blocked": record.get("is_blocked", False),
            "photo_id_status": record.get("photo_id_status", "none")}

# ---------------------------------------------------------------------------
# DEADLINE CHECK / AUTO-ESCALATION ENDPOINT
# ---------------------------------------------------------------------------
@app.post("/grievances/check-deadlines")
async def check_deadlines(user=Depends(require_role(UserRole.OFFICER.value, UserRole.ADMIN.value)),
                          db=Depends(get_db)):
    loop = asyncio.get_event_loop()
    now = datetime.now(timezone.utc)
    def find_breached():
        return list(db.grievances.find({
            "status": {"$in": ["pending", "in_progress"]},
            "$or": [
                {"sla_deadline": {"$lt": now}},
                {"estimated_resolution_deadline": {"$lt": now, "$ne": None}}
            ]
        }).limit(100))
    breached = await loop.run_in_executor(executor, find_breached)
    escalated_ids = []
    for g in breached:
        if g["status"] != "escalated":
            note = {"id": str(uuid.uuid4()), "content": "Auto-escalated: deadline breached",
                    "officer": "System", "note_type": "internal",
                    "created_at": now}
            update_fields: dict = {"status": "escalated", "updated_at": now}
            if not g.get("assigned_officer"):
                officer_name = find_officer_for_department(db, g.get("department"))
                if officer_name:
                    update_fields["assigned_officer"] = officer_name
            await loop.run_in_executor(executor, lambda gid=g["_id"], uf=update_fields: db.grievances.update_one(
                {"_id": gid},
                {"$set": uf,
                 "$push": {"notes": note}}))
            escalated_ids.append(g["_id"])
    return {"escalated_count": len(escalated_ids), "escalated_ids": escalated_ids}

# ---------------------------------------------------------------------------
# SUB-TASK ENDPOINTS (Cross-Department Dependencies)
# ---------------------------------------------------------------------------
@app.put("/grievances/{grievance_id}/sub-tasks/{sub_task_id}/status")
async def update_sub_task_status(grievance_id: str, sub_task_id: str,
                                  new_status: str = Query(...),
                                  user=Depends(require_role(UserRole.OFFICER.value, UserRole.ADMIN.value)),
                                  db=Depends(get_db)):
    grievance_id = validate_uuid(grievance_id, "grievance_id")
    loop = asyncio.get_event_loop()
    g = await loop.run_in_executor(executor, db.grievances.find_one, {"_id": grievance_id})
    if not g:
        raise HTTPException(status_code=404, detail="Grievance not found")
    sub_tasks = g.get("sub_tasks", [])
    found = False
    for st in sub_tasks:
        if st["id"] == sub_task_id:
            st["status"] = new_status
            found = True
            break
    if not found:
        raise HTTPException(status_code=404, detail="Sub-task not found")
    all_resolved = all(st["status"] == "resolved" for st in sub_tasks) if sub_tasks else False
    update = {"sub_tasks": sub_tasks, "updated_at": datetime.now(timezone.utc)}
    if all_resolved:
        update["status"] = GrievanceStatus.RESOLVED.value
    await loop.run_in_executor(executor, lambda: db.grievances.update_one(
        {"_id": grievance_id}, {"$set": update}))
    return {"detail": "Sub-task updated", "all_resolved": all_resolved}

@app.put("/grievances/{grievance_id}/sub-tasks/{sub_task_id}/assign")
async def assign_sub_task(grievance_id: str, sub_task_id: str,
                          assignment: GrievanceAssignment,
                          user=Depends(require_role(UserRole.OFFICER.value, UserRole.ADMIN.value)),
                          db=Depends(get_db)):
    grievance_id = validate_uuid(grievance_id, "grievance_id")
    loop = asyncio.get_event_loop()
    g = await loop.run_in_executor(executor, db.grievances.find_one, {"_id": grievance_id})
    if not g:
        raise HTTPException(status_code=404, detail="Grievance not found")
    sub_tasks = g.get("sub_tasks", [])
    found = False
    for st in sub_tasks:
        if st["id"] == sub_task_id:
            st["assigned_officer"] = assignment.officer_name
            st["status"] = "in_progress"
            found = True
            break
    if not found:
        raise HTTPException(status_code=404, detail="Sub-task not found")
    await loop.run_in_executor(executor, lambda: db.grievances.update_one(
        {"_id": grievance_id}, {"$set": {"sub_tasks": sub_tasks, "updated_at": datetime.now(timezone.utc)}}))
    return {"detail": "Sub-task assigned"}

@app.get("/grievances/multi-department")
async def get_multi_dept_grievances(
    department: Optional[str] = None,
    user=Depends(require_role(UserRole.OFFICER.value, UserRole.ADMIN.value)),
    db=Depends(get_db)):
    loop = asyncio.get_event_loop()
    fq = {"sub_tasks": {"$exists": True, "$ne": []}}
    if department:
        fq["sub_tasks.department"] = department
    grievances = await loop.run_in_executor(executor, lambda: [
        convert_db_grievance(g) for g in db.grievances.find(fq).sort("created_at", -1).limit(50)])
    return [GrievanceResponse(**g, id=g["_id"]) for g in grievances]

# ---------------------------------------------------------------------------
# SYSTEMIC ISSUES ENDPOINTS
# ---------------------------------------------------------------------------
@app.get("/systemic-issues")
async def get_systemic_issues(
    department: Optional[str] = None, district: Optional[str] = None,
    issue_status: Optional[str] = Query(None, alias="status"),
    user=Depends(require_role(UserRole.OFFICER.value, UserRole.ADMIN.value)),
    db=Depends(get_db)):
    fq = {}
    if department: fq["department"] = department
    if district: fq["district"] = district
    if issue_status: fq["status"] = issue_status
    loop = asyncio.get_event_loop()
    issues = await loop.run_in_executor(executor, lambda: list(
        db.systemic_issues.find(fq).sort("created_at", -1).limit(50)))
    return [SystemicIssueResponse(**{**i, "id": i["_id"]}) for i in issues]

@app.get("/systemic-issues/{issue_id}")
async def get_systemic_issue(issue_id: str,
                             user=Depends(require_role(UserRole.OFFICER.value, UserRole.ADMIN.value)),
                             db=Depends(get_db)):
    loop = asyncio.get_event_loop()
    issue = await loop.run_in_executor(executor, db.systemic_issues.find_one, {"_id": issue_id})
    if not issue:
        raise HTTPException(status_code=404, detail="Systemic issue not found")
    return SystemicIssueResponse(**{**issue, "id": issue["_id"]})

@app.put("/systemic-issues/{issue_id}/status")
async def update_systemic_issue_status(issue_id: str, new_status: str = Query(...),
                                        user=Depends(require_role(UserRole.OFFICER.value, UserRole.ADMIN.value)),
                                        db=Depends(get_db)):
    loop = asyncio.get_event_loop()
    result = await loop.run_in_executor(executor, lambda: db.systemic_issues.update_one(
        {"_id": issue_id}, {"$set": {"status": new_status, "updated_at": datetime.now(timezone.utc)}}))
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Systemic issue not found")
    return {"detail": "Status updated"}

@app.put("/systemic-issues/{issue_id}/assign")
async def assign_systemic_issue(issue_id: str, assignment: GrievanceAssignment,
                                 user=Depends(require_role(UserRole.OFFICER.value, UserRole.ADMIN.value)),
                                 db=Depends(get_db)):
    loop = asyncio.get_event_loop()
    result = await loop.run_in_executor(executor, lambda: db.systemic_issues.update_one(
        {"_id": issue_id}, {"$set": {"assigned_officer": assignment.officer_name,
                                      "status": "acknowledged", "updated_at": datetime.now(timezone.utc)}}))
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Systemic issue not found")
    return {"detail": "Officer assigned"}

# ---------------------------------------------------------------------------
# FORECASTING ENDPOINTS
# ---------------------------------------------------------------------------
@app.post("/forecasts/generate")
async def generate_forecast(user=Depends(require_role(UserRole.ADMIN.value)), db=Depends(get_db)):
    forecast = await PredictiveForecaster.generate_forecast()
    return ForecastResponse(**{**forecast, "id": forecast["_id"]})

@app.get("/forecasts/latest")
async def get_latest_forecast(user=Depends(require_role(UserRole.OFFICER.value, UserRole.ADMIN.value)),
                              db=Depends(get_db)):
    loop = asyncio.get_event_loop()
    f = await loop.run_in_executor(executor, lambda: db.forecasts.find_one(sort=[("forecast_date", -1)]))
    if not f:
        return {"detail": "No forecasts available"}
    return ForecastResponse(**{**f, "id": f["_id"]})

@app.get("/forecasts")
async def get_forecasts(limit: int = Query(10, ge=1, le=50),
                        user=Depends(require_role(UserRole.OFFICER.value, UserRole.ADMIN.value)),
                        db=Depends(get_db)):
    loop = asyncio.get_event_loop()
    forecasts = await loop.run_in_executor(executor, lambda: list(
        db.forecasts.find().sort("forecast_date", -1).limit(limit)))
    return [ForecastResponse(**{**f, "id": f["_id"]}) for f in forecasts]

# ---------------------------------------------------------------------------
# SCHEME-MATCHED GRIEVANCES ENDPOINT
# ---------------------------------------------------------------------------
@app.get("/grievances/scheme-matched")
async def get_scheme_matched_grievances(
    user=Depends(require_role(UserRole.OFFICER.value, UserRole.ADMIN.value)),
    db=Depends(get_db)):
    loop = asyncio.get_event_loop()
    grievances = await loop.run_in_executor(executor, lambda: [
        convert_db_grievance(g) for g in db.grievances.find({
            "scheme_match": {"$exists": True, "$ne": None},
            "status": {"$in": ["pending", "in_progress", "escalated"]}
        }).sort("created_at", -1).limit(50)])
    return [GrievanceResponse(**g, id=g["_id"]) for g in grievances]

# ---------------------------------------------------------------------------
# OFFICER ANOMALY ENDPOINTS
# ---------------------------------------------------------------------------
@app.get("/admin/officer-anomalies")
async def get_officer_anomalies(user=Depends(require_role(UserRole.ADMIN.value)), db=Depends(get_db)):
    loop = asyncio.get_event_loop()
    records = await loop.run_in_executor(executor, lambda: list(
        db.officer_analytics.find({"anomaly_flags": {"$exists": True, "$ne": []}}).limit(50)))
    result = []
    for r in records:
        u = await loop.run_in_executor(executor, db.users.find_one, {"_id": r["_id"]})
        result.append({
            "officer_id": r["_id"],
            "username": u["username"] if u else "unknown",
            "full_name": u["full_name"] if u else "unknown",
            "anomaly_flags": r.get("anomaly_flags", []),
            "daily_resolution_counts": r.get("daily_resolution_counts", {}),
            "priority_distribution": r.get("priority_distribution", {}),
        })
    return result

@app.get("/admin/officer-anomalies/{officer_id}")
async def get_officer_anomaly_detail(officer_id: str,
                                      user=Depends(require_role(UserRole.ADMIN.value)), db=Depends(get_db)):
    loop = asyncio.get_event_loop()
    record = await loop.run_in_executor(executor, db.officer_analytics.find_one, {"_id": officer_id})
    if not record:
        raise HTTPException(status_code=404, detail="No analytics for this officer")
    u = await loop.run_in_executor(executor, db.users.find_one, {"_id": officer_id})
    return {**record, "username": u["username"] if u else "unknown",
            "full_name": u["full_name"] if u else "unknown"}

@app.put("/admin/officer-anomalies/{officer_id}/dismiss/{flag_id}")
async def dismiss_anomaly_flag(officer_id: str, flag_id: str,
                               user=Depends(require_role(UserRole.ADMIN.value)), db=Depends(get_db)):
    loop = asyncio.get_event_loop()
    await loop.run_in_executor(executor, lambda: db.officer_analytics.update_one(
        {"_id": officer_id}, {"$pull": {"anomaly_flags": {"id": flag_id}}}))
    return {"detail": "Flag dismissed"}

# ---------------------------------------------------------------------------
# PUBLIC SOCIAL FEED ENDPOINTS
# ---------------------------------------------------------------------------
@app.get("/public/grievances")
async def get_public_grievances(
    lat: Optional[float] = None, lon: Optional[float] = None,
    radius_km: float = Query(50, ge=1, le=500),
    limit: int = Query(20, ge=1, le=100), skip: int = Query(0, ge=0),
    db=Depends(get_db)):
    loop = asyncio.get_event_loop()
    if lat is not None and lon is not None:
        def fetch_near():
            return list(db.grievances.find({
                "is_public": True,
                "status": {"$ne": "resolved"},
                "location": {
                    "$near": {
                        "$geometry": {"type": "Point", "coordinates": [lon, lat]},
                        "$maxDistance": radius_km * 1000
                    }
                }
            }).skip(skip).limit(limit))
        try:
            grievances = await loop.run_in_executor(executor, fetch_near)
        except Exception:
            grievances = await loop.run_in_executor(executor, lambda: list(
                db.grievances.find({"is_public": True, "status": {"$ne": "resolved"}})
                .sort("created_at", -1).skip(skip).limit(limit)))
    else:
        grievances = await loop.run_in_executor(executor, lambda: list(
            db.grievances.find({"is_public": True, "status": {"$ne": "resolved"}})
            .sort("created_at", -1).skip(skip).limit(limit)))
    result = []
    for g in grievances:
        vouch_count = await loop.run_in_executor(executor,
            lambda gid=g["_id"]: db.vouches.count_documents({"grievance_id": gid}))
        evidence_count = await loop.run_in_executor(executor,
            lambda gid=g["_id"]: db.vouches.count_documents({"grievance_id": gid, "evidence_file_ids": {"$ne": []}}))
        dist_km = None
        if lat is not None and lon is not None and g.get("location"):
            gc = g["location"].get("coordinates", [])
            if len(gc) == 2:
                dlat = math.radians(gc[1] - lat)
                dlon = math.radians(gc[0] - lon)
                a = math.sin(dlat/2)**2 + math.cos(math.radians(lat)) * math.cos(math.radians(gc[1])) * math.sin(dlon/2)**2
                dist_km = round(6371 * 2 * math.atan2(math.sqrt(a), math.sqrt(1-a)), 1)
        result.append(PublicGrievanceResponse(
            id=g["_id"], title=g["title"],
            description_snippet=g["description"][:200] + ("..." if len(g["description"]) > 200 else ""),
            department=g.get("department", "general"), priority=g.get("priority", "medium"),
            district=g.get("district"), location=g.get("location"),
            created_at=g["created_at"], vouch_count=vouch_count,
            evidence_count=evidence_count, distance_km=dist_km))
    return result

@app.get("/public/grievances/{grievance_id}")
async def get_public_grievance_detail(grievance_id: str, db=Depends(get_db)):
    loop = asyncio.get_event_loop()
    g = await loop.run_in_executor(executor, db.grievances.find_one,
                                   {"_id": grievance_id, "is_public": True})
    if not g:
        raise HTTPException(status_code=404, detail="Public grievance not found")
    vouches = await loop.run_in_executor(executor, lambda: list(
        db.vouches.find({"grievance_id": grievance_id}).sort("created_at", -1).limit(50)))
    for v in vouches:
        v["id"] = str(v.pop("_id"))
    g = convert_db_grievance(g)
    return {"grievance": GrievanceResponse(**g, id=g["_id"]), "vouches": vouches}

@app.post("/public/grievances/{grievance_id}/vouch")
async def vouch_for_grievance(
    grievance_id: str,
    request: Request,
    background_tasks: BackgroundTasks,
    user=Depends(get_current_user),
    db=Depends(get_db),
    comment: Optional[str] = Form(None),
    files: List[UploadFile] = File(default=[]),
):
    loop = asyncio.get_event_loop()
    g = await loop.run_in_executor(executor, db.grievances.find_one,
                                   {"_id": grievance_id, "is_public": True})
    if not g:
        raise HTTPException(status_code=404, detail="Public grievance not found")
    if g.get("citizen_user_id") == str(user["_id"]):
        raise HTTPException(status_code=400, detail="Cannot vouch for your own grievance")
    existing = await loop.run_in_executor(executor, lambda: db.vouches.find_one(
        {"grievance_id": grievance_id, "user_id": str(user["_id"])}))
    if existing:
        raise HTTPException(status_code=400, detail="You have already vouched for this grievance")
    evidence_ids = []
    for f in files[:3]:
        if f.content_type not in ALLOWED_MIME_TYPES:
            continue
        content = await f.read()
        if len(content) > MAX_FILE_SIZE:
            continue
        fid = await loop.run_in_executor(executor, lambda c=content, fn=f.filename, ct=f.content_type: gfs.put(
            c, filename=fn, content_type=ct))
        evidence_ids.append(str(fid))
    vouch_doc = {
        "_id": str(uuid.uuid4()), "grievance_id": grievance_id,
        "user_id": str(user["_id"]),
        "comment": comment[:500] if comment else None,
        "evidence_file_ids": evidence_ids,
        "created_at": datetime.now(timezone.utc),
    }
    await loop.run_in_executor(executor, db.vouches.insert_one, vouch_doc)
    return {"detail": "Vouch recorded", "vouch_id": vouch_doc["_id"]}

@app.get("/public/grievances/{grievance_id}/vouches")
async def get_vouches(grievance_id: str, db=Depends(get_db)):
    loop = asyncio.get_event_loop()
    vouches = await loop.run_in_executor(executor, lambda: list(
        db.vouches.find({"grievance_id": grievance_id}).sort("created_at", -1).limit(50)))
    for v in vouches:
        v["id"] = str(v.pop("_id"))
    return vouches

# ---------------------------------------------------------------------------
# HEALTH
# ---------------------------------------------------------------------------
@app.get("/health")
async def health():
    return {"status": "healthy", "system": "PR&DW Grievance Portal",
            "timestamp": datetime.now(timezone.utc)}

# ---------------------------------------------------------------------------
# PAGE ROUTES (serve Jinja2 templates)
# ---------------------------------------------------------------------------
PAGES = [
    ("", "login.html"), ("login", "login.html"), ("register", "register.html"),
    ("dashboard", "dashboard.html"), ("file-grievance", "file_grievance.html"),
    ("track", "track.html"), ("chatbot", "chatbot.html"), ("schemes", "schemes.html"),
    ("officer-dashboard", "officer_dashboard.html"), ("queue", "queue.html"),
    ("grievance-detail", "grievance_detail.html"), ("knowledge", "knowledge.html"),
    ("scheme-detail", "scheme_detail.html"),
    ("analytics-view", "analytics.html"),
    ("admin", "admin.html"),
    ("community", "community.html"),
    ("systemic-issue-detail", "systemic_issue_detail.html"),
    ("spam-flag-detail", "spam_flag_detail.html"),
]

for _path, _template in PAGES:
    def _make_handler(tmpl: str):
        async def handler(request: Request):
            response = templates.TemplateResponse(tmpl, {"request": request})
            response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
            response.headers["Pragma"] = "no-cache"
            return response
        return handler
    app.add_api_route(f"/{_path}" if _path else "/", _make_handler(_template),
                      methods=["GET"], response_class=HTMLResponse, include_in_schema=False)

# ---------------------------------------------------------------------------
# ADMIN REPORTS ENDPOINTS
# ---------------------------------------------------------------------------
@app.get("/admin/reports", response_class=HTMLResponse)
async def admin_reports_page(request: Request):
    return templates.TemplateResponse("reports.html", {"request": request})

@app.get("/admin/reports/stats")
async def admin_reports_stats(
    timeframe: str = Query("daily", regex="^(daily|weekly)$"),
    user=Depends(require_role(UserRole.ADMIN.value)),
    db=Depends(get_db)
):
    """
    Get statistics for Admin Reports.
    - Daily: usage since midnight local time (approx via UTC offset or just UTC midnight).
      For simplicity in this demo, 'daily' = since 00:00 UTC today.
    - Weekly: usage since 7 days ago.
    """
    now = datetime.now(timezone.utc)
    if timeframe == "daily":
        # Start of today (UTC)
        start_date = now.replace(hour=0, minute=0, second=0, microsecond=0)
    else:
        # Last 7 days
        start_date = now - timedelta(days=7)

    def fetch_report_data():
        # 1. Total Active Grievances (Snapshot, independent of timeframe)
        total_active = db.grievances.count_documents({"status": {"$in": ["pending", "in_progress", "escalated"]}})

        # 2. New Grievances Added in Timeframe
        new_grievances_cursor = db.grievances.find({"created_at": {"$gte": start_date}}).sort("created_at", -1)
        new_grievances_list = list(new_grievances_cursor)
        new_count = len(new_grievances_list)
        
        # 3. Resolved in Timeframe
        resolved_count = db.grievances.count_documents({
            "updated_at": {"$gte": start_date},
            "status": "resolved"
        })
        
        # 4. Status Breakdown (filtered by timeframe)
        # User requested Daily/Weekly distribution
        pipeline = [
            {"$match": {"created_at": {"$gte": start_date}}},
            {"$group": {"_id": "$status", "count": {"$sum": 1}}}
        ]
        status_results = list(db.grievances.aggregate(pipeline))
        status_counts = {r["_id"]: r["count"] for r in status_results}
        print(f"DEBUG: Status Breakdown Results ({timeframe}): {status_results}") # Debug logging

        # District-wise (New)
        district_counts = {}
        # Department-wise (New)
        dept_counts = {}
        # Assigned-to
        officer_counts = {}

        converted_new_list = []
        for g in new_grievances_list:
            d = g.get("district") or "Unknown"
            district_counts[d] = district_counts.get(d, 0) + 1
            
            dept = g.get("department") or "general"
            dept_counts[dept] = dept_counts.get(dept, 0) + 1
            
            off = g.get("assigned_officer")
            if off:
                officer_counts[off] = officer_counts.get(off, 0) + 1
            
            converted_new_list.append(convert_db_grievance(g))

        # 5. Top Officers by Active Assignment (Snapshot of backlog)
        active_officer_pipe = [
            {"$match": {"status": {"$in": ["in_progress", "escalated"]}, "assigned_officer": {"$exists": True, "$ne": None}}},
            {"$group": {"_id": "$assigned_officer", "count": {"$sum": 1}}},
            {"$sort": {"count": -1}},
            {"$limit": 10}
        ]
        active_officer_results = list(db.grievances.aggregate(active_officer_pipe))
        active_officer_load = [{"name": r["_id"], "count": r["count"]} for r in active_officer_results]

        return {
            "period": "Today" if timeframe == "daily" else "Last 7 Days",
            "kpi": {
                "active_total": total_active,
                "new_added": new_count,
                "resolved": resolved_count
            },
            "status_breakdown": [{"name": k, "count": v} for k,v in status_counts.items()],
            "districts": [{"name": k, "count": v} for k,v in sorted(district_counts.items(), key=lambda item: item[1], reverse=True)],
            "departments": [{"name": k, "count": v} for k,v in sorted(dept_counts.items(), key=lambda item: item[1], reverse=True)],
            "new_by_officer": [{"name": k, "count": v} for k,v in sorted(officer_counts.items(), key=lambda item: item[1], reverse=True)],
            "current_officer_load": active_officer_load,
            "recent_grievances": [GrievanceResponse(**g, id=g["_id"]) for g in converted_new_list[:50]]
        }

    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(executor, fetch_report_data)

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)
